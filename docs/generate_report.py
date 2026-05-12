#!/usr/bin/env python3
"""Generate Phase II Design Report as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_header_footer():
    """Add header/footer to all sections"""
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        run1 = hp.add_run("CunyZeroLite")
        run1.font.size = Pt(9)
        run1.font.name = 'Times New Roman'
        hp.add_run("\t\t").font.size = Pt(9)
        run2 = hp.add_run("Version: 1.0")
        run2.font.size = Pt(9)
        run2.font.name = 'Times New Roman'

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Confidential              \u00a9 Group K")
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    return table


# ════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Group K")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

doc.add_paragraph()  # spacer

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CunyZeroLite\nDesign Report\nPhase II")
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ════════════════════════════════════════
# REVISION HISTORY
# ════════════════════════════════════════
heading("Revision History", level=1)
add_table(
    ["Date", "Version", "Description", "Author"],
    [["21/04/26", "1.0",
      "Phase II Design Report covering system architecture, all seventeen use cases with scenarios, "
      "collaboration/sequence diagrams, Petri-nets, E-R diagram, detailed pseudo-code, and system screens.",
      "Diego Reyes Liranzo, Daniel Olekszyk, Samia Islam, Maisha Islam, Kyle Gosine"]]
)

doc.add_page_break()

# ════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════
heading("Table of Contents", level=1)
toc_items = [
    "1.  Introduction",
    "    1.1  Purpose",
    "    1.2  System Collaboration Class Diagram",
    "2.  Use Cases",
    "    2.1  Use-Case Scenarios (UC-1 through UC-17)",
    "    2.2  Collaboration and Sequence Diagrams",
    "    2.3  Petri-Net Diagrams",
    "3.  E-R Diagram",
    "4.  Detailed Design (Pseudo-Code)",
    "5.  System Screens",
    "6.  Meeting Memos",
    "7.  Repository",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════
heading("1. Introduction", level=1)

heading("1.1 Purpose", level=2)
para("This Phase II Design Report provides the data structures and logic required to implement "
     "the CUNYZeroLite college management system as specified in the Phase I Software Requirements "
     "Specification. It contains the Entity-Relationship model, detailed pseudo-code for every method, "
     "use-case scenarios with collaboration/sequence diagrams, Petri-nets, and representative system "
     "screen descriptions. The implementation will be based entirely on this document.")

heading("1.2 System Collaboration Class Diagram", level=2)
para("The CUNYZeroLite system is organized around a three-layer architecture:")

code_block(
    "+-------------------------------------------------------------------+\n"
    "|                       <<boundary>>                                 |\n"
    "|                    Browser / Client                                |\n"
    "|  [LoginPage] [Dashboard] [StudentPortal] [InstructorPortal]       |\n"
    "|  [RegistrarPortal] [PublicHomePage] [AIChatWidget]                 |\n"
    "+-------------------------------|------------------------------------+\n"
    "                                |  HTTP / Server Actions\n"
    "+-------------------------------|------------------------------------+\n"
    "|                       <<control>>                                  |\n"
    "|                    Next.js API Routes & Server Actions              |\n"
    "|  [loginUser] [logoutUser] [changePassword] [getSession]           |\n"
    "|  [registerForCourse] [assignGrade] [submitReview]                 |\n"
    "|  [fileComplaint] [processComplaint] [manageApplication]           |\n"
    "|  [manageSemester] [setupCourse] [applyGraduation]                 |\n"
    "|  [manageTabooWords] [enforceRunningRules] [askAI]                 |\n"
    "+-------------------------------|------------------------------------+\n"
    "                                |  Prisma ORM\n"
    "+-------------------------------|------------------------------------+\n"
    "|                       <<entity>>                                   |\n"
    "|                    SQLite Database (via Prisma)                     |\n"
    "|  [User] [Semester] [Course] [Enrollment] [Waitlist]               |\n"
    "|  [Review] [Complaint] [Warning] [Application]                     |\n"
    "|  [GraduationRequest] [TabooWord] [HonorRoll]                     |\n"
    "+-------------------------------------------------------------------+"
)

para("The Browser sends requests to the Next.js control layer. Server Actions (loginUser, registerForCourse, etc.) "
     "validate input, enforce business rules, and interact with the Prisma ORM entity layer. The entity layer "
     "persists data in SQLite. Session management uses HTTP-only cookies. The AI subsystem queries a local "
     "policy.json knowledge base and falls back to OpenAI GPT-4o-mini.")

doc.add_page_break()

# ════════════════════════════════════════
# 2. USE CASES
# ════════════════════════════════════════
heading("2. Use Cases", level=1)
heading("2.1 Use-Case Scenarios", level=2)

use_cases = [
    ("UC-1: View Public Homepage", [
        ("Normal Scenario:", [
            "Visitor navigates to the homepage.",
            "System displays general introduction, highest-rated courses, lowest-rated courses, and top-GPA students.",
            "Lite AI chat widget is visible in the corner.",
            "Visitor reads information or interacts with Lite.",
        ]),
        ("Exceptional Scenarios:", [
            "No data exists: System displays placeholder text ('No courses available yet').",
            "Database unreachable: System shows a generic error page.",
        ]),
    ]),
    ("UC-2: Apply as Student or Instructor", [
        ("Normal Scenario:", [
            "Visitor clicks 'Apply' on the homepage.",
            "System presents application form.",
            "Visitor selects role (Student or Instructor).",
            "For Student: enters prior GPA and written justification.",
            "For Instructor: enters name and contact info (no justification required).",
            "Visitor submits the form.",
            "System creates an Application record with status PENDING and displays confirmation.",
        ]),
        ("Exceptional Scenarios:", [
            "Missing required fields: System highlights empty fields and prevents submission.",
            "Prior GPA out of range (Student): System rejects values outside 0.0-4.0.",
            "Duplicate application: System informs the visitor that a pending application already exists.",
        ]),
    ]),
    ("UC-3: Approve/Reject Applications", [
        ("Normal Scenario:", [
            "Registrar navigates to the Applications page.",
            "System displays all PENDING applications.",
            "For a Student application with prior GPA > 3.0 and quota not reached, system flags for auto-acceptance.",
            "Registrar clicks 'Approve.'",
            "System generates university email (first initial + last name + 00@cuny.edu) and temporary password.",
            "Application status updates to ACCEPTED; new User record created with mustChangePassword = true.",
        ]),
        ("Exceptional Scenarios:", [
            "Reject a qualified student (GPA > 3.0): System requires registrar to provide written justification.",
            "Program quota reached: System alerts registrar; registrar may still override.",
            "Instructor application rejected: Status set to REJECTED with no justification required.",
        ]),
    ]),
    ("UC-4: Manage Semester Periods", [
        ("Normal Scenario:", [
            "Registrar opens Semester Management.",
            "System shows current semester and its period.",
            "Registrar clicks 'Advance Period.'",
            "System transitions: CLASS_SETUP -> REGISTRATION -> RUNNING -> GRADING -> COMPLETED.",
            "Period updates; UI reflects the new state.",
        ]),
        ("Exceptional Scenarios:", [
            "Advance to COMPLETED triggers academic evaluation: System auto-recalculates GPAs, checks honor roll, issues termination for GPA < 2.0 or double-fail, issues warnings for GPA 2.0-2.25.",
            "No current semester exists: System prompts registrar to create one first.",
            "Advance out of order: System only allows the next sequential period.",
        ]),
    ]),
    ("UC-5: Set Up Courses", [
        ("Normal Scenario:", [
            "Registrar opens Course Management during CLASS_SETUP.",
            "Registrar enters course code, name, credits, schedule, max students, and selects an instructor.",
            "System creates the Course record linked to the current semester.",
            "Course appears in the course list.",
        ]),
        ("Exceptional Scenarios:", [
            "Duplicate course code in same semester: System rejects with error.",
            "Cancel a course: Registrar clicks cancel; system sets cancelled = true and issues a warning to the instructor.",
            "Semester not in CLASS_SETUP: System disables course creation.",
        ]),
    ]),
    ("UC-6: Register for Courses", [
        ("Normal Scenario:", [
            "Student opens Course Registration during REGISTRATION period.",
            "System displays available courses for the current semester.",
            "Student selects a course.",
            "System checks: no time conflict, course not full, student not suspended/terminated, 2-4 courses.",
            "System creates Enrollment record with status ENROLLED.",
        ]),
        ("Exceptional Scenarios:", [
            "Time conflict: System rejects and displays conflicting course.",
            "Course full: System creates Waitlist entry instead and shows waitlist position.",
            "Student suspended or terminated: System denies with message.",
            "Already enrolled in 4 courses: System prevents registration.",
            "Retaking a course not previously failed with F: System blocks re-enrollment.",
        ]),
    ]),
    ("UC-7: Manage Waitlist", [
        ("Normal Scenario:", [
            "Instructor opens their course page.",
            "System displays waitlisted students in order of request.",
            "Instructor clicks 'Admit' next to a student.",
            "System changes Waitlist status to ADMITTED, creates an Enrollment record.",
            "Waitlist positions update for remaining students.",
        ]),
        ("Exceptional Scenarios:", [
            "Course at capacity and instructor admits: System allows (instructor override).",
            "Student no longer eligible (suspended since waitlisting): System warns instructor and blocks.",
            "Empty waitlist: System displays 'No students on waitlist.'",
        ]),
    ]),
    ("UC-8: Write and Rate Course Reviews", [
        ("Normal Scenario:", [
            "Student navigates to Course Reviews.",
            "System shows enrolled courses (grades not yet posted).",
            "Student selects a course, enters 1-5 star rating and optional comment.",
            "System checks comment against TabooWord list.",
            "Zero taboo words found: review is created and published.",
            "If course average rating drops below 2.0, instructor receives automatic warning.",
        ]),
        ("Exceptional Scenarios:", [
            "1-2 taboo words: System replaces words with asterisks, publishes review, issues 1 warning to student.",
            "3+ taboo words: System hides review entirely and issues 2 warnings to student.",
            "Grades already posted: System prevents review submission.",
            "Duplicate review: System rejects (unique constraint on authorId + courseId).",
        ]),
    ]),
    ("UC-9: Assign Grades", [
        ("Normal Scenario:", [
            "Instructor opens Grading page during GRADING period.",
            "System lists all enrolled students per course.",
            "Instructor assigns a letter grade (A, B, C, D, or F) to each student.",
            "System saves grades to Enrollment records.",
            "When all students graded, system marks grading complete for that course.",
        ]),
        ("Exceptional Scenarios:", [
            "Not all students graded before period ends: Instructor receives a warning.",
            "Class GPA > 3.5 or < 2.5: Registrar may question instructor and issue warning/termination.",
            "Student fails same course twice: System auto-terminates after finalization.",
            "Student cumulative GPA < 2.0: System auto-terminates.",
            "Student GPA 2.0-2.25: Warning issued; must interview registrar.",
            "Semester GPA > 3.75 or cumulative > 3.5 (after >1 semester): Honor roll; can remove one warning.",
        ]),
    ]),
    ("UC-10: File Complaint", [
        ("Normal Scenario:", [
            "Student or Instructor navigates to File Complaint.",
            "User selects target (a student or instructor), enters description.",
            "System creates Complaint record with status PENDING.",
            "Confirmation shown.",
        ]),
        ("Exceptional Scenarios:", [
            "Complaint against self: System prevents filing.",
            "Empty description: System requires text before submission.",
            "Instructor complaint requests de-registration: Recorded in description.",
        ]),
    ]),
    ("UC-11: Process Complaints", [
        ("Normal Scenario:", [
            "Registrar opens Pending Complaints.",
            "System lists all PENDING complaints with filer and target details.",
            "Registrar reviews a complaint and chooses action: issue warning, de-register student, or dismiss.",
            "System updates Complaint status to RESOLVED and records resolution.",
        ]),
        ("Exceptional Scenarios:", [
            "Instructor-filed complaint: Registrar must either punish target or warn instructor for unjustified filing.",
            "Action triggers 3rd warning: System auto-suspends student, sets fineOwed, prevents future registration.",
            "Dismiss complaint: Status set to DISMISSED, no action taken.",
        ]),
    ]),
    ("UC-12: Apply for Graduation", [
        ("Normal Scenario:", [
            "Student navigates to Graduation Application.",
            "System verifies student has completed 8 courses with passing grades.",
            "Student submits application.",
            "System creates GraduationRequest with status PENDING.",
        ]),
        ("Exceptional Scenarios:", [
            "Fewer than 8 completed courses: System warns this is premature; student may still submit.",
            "Reckless application: If registrar rejects, student receives a warning.",
        ]),
    ]),
    ("UC-13: Approve/Reject Graduation", [
        ("Normal Scenario:", [
            "Registrar opens Graduation Requests.",
            "System displays pending requests with student course history.",
            "Registrar verifies 8 courses with passing grades and no outstanding holds.",
            "Registrar approves: student's graduated flag set to true; student leaves system.",
        ]),
        ("Exceptional Scenarios:", [
            "Outstanding holds (fines, suspension): Registrar rejects; warning issued.",
            "Fewer than 8 passing courses: Registrar rejects; reckless application warning.",
        ]),
    ]),
    ("UC-14: Ask AI Assistant (Lite)", [
        ("Normal Scenario:", [
            "User clicks the Lite chat widget.",
            "User types a question.",
            "System searches policy.json for matching keywords (RAG).",
            "Match found: relevant policy passed as context to GPT-4o-mini; grounded response returned.",
            "No match: GPT-4o-mini responds from general knowledge with hallucination warning.",
        ]),
        ("Exceptional Scenarios:", [
            "OpenAI API key invalid/missing: System returns 'AI assistant is currently unavailable.'",
            "API rate limit exceeded: System shows retry message.",
            "Empty question: System prompts user to enter a question.",
        ]),
    ]),
    ("UC-15: View Role-Based Dashboard", [
        ("Normal Scenario:", [
            "User logs in successfully.",
            "System reads session cookie and determines role.",
            "System redirects to appropriate dashboard (Student/Instructor/Registrar).",
            "Dashboard data fetched from database and rendered.",
        ]),
        ("Exceptional Scenarios:", [
            "Session expired/invalid: Redirect to login.",
            "First-time login (mustChangePassword = true): Redirect to change-password page.",
            "Account terminated/suspended/fired between sessions: Login check prevents access.",
        ]),
    ]),
    ("UC-16: Manage Taboo Words", [
        ("Normal Scenario:", [
            "Registrar opens Taboo Words page.",
            "System displays current list of taboo words.",
            "Registrar adds a new word; system creates TabooWord record.",
            "Registrar removes a word; system deletes the record.",
        ]),
        ("Exceptional Scenarios:", [
            "Duplicate word: System rejects (unique constraint).",
            "Empty input: System prevents submission.",
        ]),
    ]),
    ("UC-17: Enforce Class Running Period Rules", [
        ("Normal Scenario:", [
            "Registrar advances semester to RUNNING.",
            "System evaluates: students with < 2 courses get a warning.",
            "Courses with < 3 students are cancelled; affected students get special registration window.",
            "Instructors of cancelled courses receive a warning.",
        ]),
        ("Exceptional Scenarios:", [
            "Instructor's entire course load cancelled: Instructor suspended, cannot teach next semester.",
            "Student already at 3 warnings after check: Auto-suspension triggered.",
            "No courses meet cancellation criteria: System proceeds without changes.",
        ]),
    ]),
]

for uc_title, scenarios in use_cases:
    heading(uc_title, level=3)
    for scenario_title, steps in scenarios:
        para(scenario_title, bold=True)
        for i, step in enumerate(steps, 1):
            if "Exceptional" in scenario_title:
                bullet(step)
            else:
                bullet(f"{i}. {step}")

doc.add_page_break()

# ════════════════════════════════════════
# 2.2 SEQUENCE / COLLABORATION DIAGRAMS
# ════════════════════════════════════════
heading("2.2 Collaboration and Sequence Diagrams", level=2)

heading("UC-1: View Public Homepage (Sequence Diagram)", level=3)
code_block(
    "Visitor          Browser          Server           Database\n"
    "  |                |                |                  |\n"
    "  |-- open URL --->|                |                  |\n"
    "  |                |-- GET / ------>|                  |\n"
    "  |                |                |-- query top      |\n"
    "  |                |                |   courses,       |\n"
    "  |                |                |   students ----->|\n"
    "  |                |                |<-- result set ---|\n"
    "  |                |<-- HTML page --|                  |\n"
    "  |<-- display ----|                |                  |"
)

heading("UC-2: Apply as Student or Instructor (Sequence Diagram)", level=3)
code_block(
    "Visitor          Browser           Server            Database\n"
    "  |                |                 |                  |\n"
    "  |-- fill form -->|                 |                  |\n"
    "  |                |-- POST /api/    |                  |\n"
    "  |                |   apply ------->|                  |\n"
    "  |                |                 |-- validate       |\n"
    "  |                |                 |   fields         |\n"
    "  |                |                 |-- check dup ---->|\n"
    "  |                |                 |<-- no dup -------|\n"
    "  |                |                 |-- INSERT         |\n"
    "  |                |                 |   Application -->|\n"
    "  |                |                 |<-- OK -----------|\n"
    "  |                |<-- 200 confirm -|                  |\n"
    "  |<-- show msg ---|                 |                  |"
)

heading("UC-6: Register for Courses (Sequence Diagram)", level=3)
code_block(
    "Student          Browser           Server            Database\n"
    "  |                |                 |                  |\n"
    "  |-- select       |                 |                  |\n"
    "  |   course ----->|                 |                  |\n"
    "  |                |-- POST          |                  |\n"
    "  |                |   /register --->|                  |\n"
    "  |                |                 |-- getSession()   |\n"
    "  |                |                 |-- check          |\n"
    "  |                |                 |   suspended? --->|\n"
    "  |                |                 |<-- no -----------|\n"
    "  |                |                 |-- check time     |\n"
    "  |                |                 |   conflict ----->|\n"
    "  |                |                 |<-- no conflict --|\n"
    "  |                |                 |-- check capacity>|\n"
    "  |                |                 |<-- spots open ---|\n"
    "  |                |                 |-- INSERT         |\n"
    "  |                |                 |   Enrollment --->|\n"
    "  |                |                 |<-- OK -----------|\n"
    "  |                |<-- success -----|                  |\n"
    "  |<-- updated UI -|                 |                  |"
)

heading("UC-9: Assign Grades (Sequence Diagram)", level=3)
code_block(
    "Instructor       Browser           Server            Database\n"
    "  |                |                 |                  |\n"
    "  |-- enter        |                 |                  |\n"
    "  |   grades ----->|                 |                  |\n"
    "  |                |-- POST          |                  |\n"
    "  |                |   /grades ----->|                  |\n"
    "  |                |                 |-- getSession()   |\n"
    "  |                |                 |-- validate all   |\n"
    "  |                |                 |   graded?        |\n"
    "  |                |                 |-- UPDATE grades->|\n"
    "  |                |                 |<-- OK -----------|\n"
    "  |                |                 |-- recalc GPA     |\n"
    "  |                |                 |-- check honors   |\n"
    "  |                |                 |-- check term --->|\n"
    "  |                |                 |<-- done ---------|\n"
    "  |                |<-- success -----|                  |\n"
    "  |<-- confirmed --|                 |                  |"
)

heading("UC-15: View Role-Based Dashboard (Collaboration Diagram)", level=3)
code_block(
    "              1: request dashboard\n"
    "[Browser] --------------------------------> [DashboardController]\n"
    "   ^                                              |\n"
    "   |                                    2: getSession()\n"
    "   |                                              v\n"
    "   |                                     [SessionManager]\n"
    "   |                                    3: read cookie\n"
    "   |                              4: determine role\n"
    "   |                              5: redirect to role dashboard\n"
    "   |                                              v\n"
    "   |                                  [RoleDashboard]\n"
    "   |                             6: query user data,\n"
    "   |                                enrollments, warnings\n"
    "   |                                              v\n"
    "   |                                         [Database]\n"
    "   |                             7: return data\n"
    "   |           8: render HTML                     |\n"
    "   |<---------------------------------------------|"
)

heading("UC-11: Process Complaints (Collaboration Diagram)", level=3)
code_block(
    "              1: open complaints\n"
    "[Browser] --------------------------------> [ComplaintController]\n"
    "   ^                                              |\n"
    "   |                                    2: getSession()\n"
    "   |                                    3: verify REGISTRAR\n"
    "   |                              4: fetch PENDING complaints\n"
    "   |                                              v\n"
    "   |                                         [Database]\n"
    "   |           5: display list                    |\n"
    "   |<---------------------------------------------|\n"
    "   |                                              |\n"
    "   |-- 6: select action (warn/dismiss/punish) --->|\n"
    "   |                              7: UPDATE complaint\n"
    "   |                              8: IF warn: INSERT Warning\n"
    "   |                              9: IF 3 warnings: suspend\n"
    "   |          10: confirmation                    |\n"
    "   |<---------------------------------------------|"
)

doc.add_page_break()

# ════════════════════════════════════════
# 2.3 PETRI-NET DIAGRAMS
# ════════════════════════════════════════
heading("2.3 Petri-Net Diagrams", level=2)

heading("Petri-Net 1: UC-6 \u2014 Register for Courses", level=3)
code_block(
    "                        (P1)\n"
    "                     [Student at\n"
    "                   Registration Page]\n"
    "                         |\n"
    "                         | t1: select course\n"
    "                         v\n"
    "                        (P2)\n"
    "                    [Course Selected]\n"
    "                         |\n"
    "            +------------+-----------+\n"
    "            |                        |\n"
    "     t2: eligible             t3: not eligible\n"
    "        (PASS)                (suspended/terminated)\n"
    "            |                        |\n"
    "            v                        v\n"
    "           (P3)                   (P_DENIED)\n"
    "     [Eligibility OK]          [Registration Denied]\n"
    "            |\n"
    "     +------+------+\n"
    "     |             |\n"
    "  t4: no         t5: time\n"
    "  conflict       conflict\n"
    "     |              |\n"
    "     v              v\n"
    "    (P4)         (P_CONFLICT)\n"
    "  [No Conflict]  [Conflict Error]\n"
    "     |\n"
    "     +------+------+\n"
    "     |             |\n"
    "  t6: has        t7: course\n"
    "  spots          is full\n"
    "     |              |\n"
    "     v              v\n"
    "    (P5)          (P6)\n"
    " [Spot Open]   [Course Full]\n"
    "     |              |\n"
    "  t8: INSERT     t9: INSERT\n"
    "  Enrollment     Waitlist\n"
    "     |              |\n"
    "     v              v\n"
    "    (P7)          (P8)\n"
    " [Enrolled]    [Waitlisted]"
)

heading("Petri-Net 2: UC-9 \u2014 Assign Grades (Academic Standing)", level=3)
code_block(
    "                     (P1)\n"
    "                 [Grading Period Active]\n"
    "                      |\n"
    "                      | t1: instructor opens grading\n"
    "                      v\n"
    "                     (P2)\n"
    "                 [Student List Displayed]\n"
    "                      |\n"
    "                      | t2: assign letter grade\n"
    "                      v\n"
    "                     (P3)\n"
    "                 [Grade Saved]\n"
    "                      |\n"
    "            +---------+---------+\n"
    "            |                   |\n"
    "         t3: all             t4: not all\n"
    "         graded              graded (period ends)\n"
    "            |                   |\n"
    "            v                   v\n"
    "           (P4)              (P_WARN)\n"
    "       [Grades             [Instructor Warning]\n"
    "        Finalized]\n"
    "            |\n"
    "            | t5: recalculate GPA\n"
    "            v\n"
    "           (P5)\n"
    "       [GPA Updated]\n"
    "            |\n"
    "    +-------+--------+--------+\n"
    "    |                |         |\n"
    " t6: GPA<2.0     t7: GPA    t8: GPA>=2.25\n"
    " or double-fail  2.0-2.25       |\n"
    "    |                |     +----+-----+\n"
    "    v                v     |          |\n"
    " (P_TERM)     (P_PROBATION) t9:honor  t10:normal\n"
    " [Terminated]  [Warning+    |          |\n"
    "               Interview]   v          v\n"
    "                        (P_HONOR)  (P_DONE)\n"
    "                        [Honor     [Complete]\n"
    "                         Roll]"
)

heading("Petri-Net 3: UC-8 \u2014 Write Course Reviews (Taboo Filtering)", level=3)
code_block(
    "                     (P1)\n"
    "                 [Student on Reviews Page]\n"
    "                      |\n"
    "                      | t1: enter rating + comment\n"
    "                      v\n"
    "                     (P2)\n"
    "                 [Review Submitted]\n"
    "                      |\n"
    "                      | t2: check TabooWord list\n"
    "                      v\n"
    "                     (P3)\n"
    "                 [Taboo Count Determined]\n"
    "                      |\n"
    "        +-------------+-------------+\n"
    "        |             |              |\n"
    "     t3: 0          t4: 1-2       t5: 3+\n"
    "     taboo          taboo         taboo\n"
    "        |             |              |\n"
    "        v             v              v\n"
    "       (P4)         (P5)           (P7)\n"
    "    [Published]   [Words ***'d,  [Hidden,\n"
    "        |          1 Warning]     2 Warnings]\n"
    "        |             |              |\n"
    "        v             v              v\n"
    "       (P8)         (P8)           (P8)\n"
    "    [Check Avg    [Check 3-warn  [Check 3-warn\n"
    "     Rating]       threshold]     threshold]\n"
    "        |\n"
    "        | t6: avg < 2.0\n"
    "        v\n"
    "       (P9)\n"
    "    [Instructor Warning]"
)

doc.add_page_break()

# ════════════════════════════════════════
# 3. E-R DIAGRAM
# ════════════════════════════════════════
heading("3. E-R Diagram", level=1)
para("The CUNYZeroLite database consists of 11 entities. Below are all entities with their attributes, "
     "primary keys (PK), foreign keys (FK), and constraints.")

entities = [
    ("User", [
        ["id", "Integer", "PK, auto-increment"],
        ["email", "String", "UNIQUE"],
        ["username", "String", "UNIQUE"],
        ["password", "String", ""],
        ["firstName", "String", ""],
        ["lastName", "String", ""],
        ["role", "String", "Default: STUDENT"],
        ["gpa", "Float", "Default: 0.0"],
        ["warnings", "Int", "Default: 0"],
        ["suspended", "Boolean", "Default: false"],
        ["terminated", "Boolean", "Default: false"],
        ["fired", "Boolean", "Default: false"],
        ["graduated", "Boolean", "Default: false"],
        ["fineOwed", "Float", "Default: 0.0"],
        ["mustChangePassword", "Boolean", "Default: true"],
        ["createdAt", "DateTime", "Auto"],
        ["updatedAt", "DateTime", "Auto"],
    ]),
    ("Semester", [
        ["id", "Integer", "PK, auto-increment"],
        ["name", "String", "UNIQUE"],
        ["year", "Int", ""],
        ["term", "String", ""],
        ["period", "String", "Default: CLASS_SETUP"],
        ["startDate", "DateTime", ""],
        ["endDate", "DateTime", ""],
        ["isCurrent", "Boolean", "Default: false"],
        ["programQuota", "Int", "Default: 50"],
        ["createdAt", "DateTime", "Auto"],
    ]),
    ("Course", [
        ["id", "Integer", "PK, auto-increment"],
        ["code", "String", "UNIQUE(code, semesterId)"],
        ["name", "String", ""],
        ["credits", "Int", "Default: 3"],
        ["maxStudents", "Int", "Default: 30"],
        ["schedule", "String", ""],
        ["cancelled", "Boolean", "Default: false"],
        ["createdAt", "DateTime", "Auto"],
        ["semesterId", "Int", "FK -> Semester.id"],
        ["instructorId", "Int?", "FK -> User.id (nullable)"],
    ]),
    ("Enrollment", [
        ["id", "Integer", "PK, auto-increment"],
        ["status", "String", "Default: ENROLLED"],
        ["grade", "String?", "Nullable"],
        ["createdAt", "DateTime", "Auto"],
        ["userId", "Int", "FK -> User.id, UNIQUE(userId,courseId)"],
        ["courseId", "Int", "FK -> Course.id"],
    ]),
    ("Waitlist", [
        ["id", "Integer", "PK, auto-increment"],
        ["status", "String", "Default: WAITING"],
        ["position", "Int", ""],
        ["createdAt", "DateTime", "Auto"],
        ["userId", "Int", "FK -> User.id, UNIQUE(userId,courseId)"],
        ["courseId", "Int", "FK -> Course.id"],
    ]),
    ("Review", [
        ["id", "Integer", "PK, auto-increment"],
        ["rating", "Int", "1-5"],
        ["comment", "String?", "Nullable"],
        ["hidden", "Boolean", "Default: false"],
        ["createdAt", "DateTime", "Auto"],
        ["authorId", "Int", "FK -> User.id, UNIQUE(authorId,courseId)"],
        ["courseId", "Int", "FK -> Course.id"],
    ]),
    ("Complaint", [
        ["id", "Integer", "PK, auto-increment"],
        ["description", "String", ""],
        ["status", "String", "Default: PENDING"],
        ["resolution", "String?", "Nullable"],
        ["createdAt", "DateTime", "Auto"],
        ["filerId", "Int", "FK -> User.id"],
        ["targetId", "Int", "FK -> User.id"],
    ]),
    ("Warning", [
        ["id", "Integer", "PK, auto-increment"],
        ["reason", "String", ""],
        ["removed", "Boolean", "Default: false"],
        ["createdAt", "DateTime", "Auto"],
        ["userId", "Int", "FK -> User.id"],
    ]),
    ("Application", [
        ["id", "Integer", "PK, auto-increment"],
        ["type", "String", "STUDENT | INSTRUCTOR"],
        ["status", "String", "Default: PENDING"],
        ["priorGpa", "Float?", "Nullable"],
        ["justification", "String?", "Nullable"],
        ["createdAt", "DateTime", "Auto"],
        ["userId", "Int", "FK -> User.id"],
    ]),
    ("GraduationRequest", [
        ["id", "Integer", "PK, auto-increment"],
        ["status", "String", "Default: PENDING"],
        ["createdAt", "DateTime", "Auto"],
        ["userId", "Int", "FK -> User.id"],
    ]),
    ("TabooWord", [
        ["id", "Integer", "PK, auto-increment"],
        ["word", "String", "UNIQUE"],
    ]),
    ("HonorRoll", [
        ["id", "Integer", "PK, auto-increment"],
        ["type", "String", "SEMESTER | OVERALL"],
        ["usedToRemoveWarning", "Boolean", "Default: false"],
        ["createdAt", "DateTime", "Auto"],
        ["userId", "Int", "FK -> User.id, UNIQUE(userId,semesterId,type)"],
        ["semesterId", "Int", "FK -> Semester.id"],
    ]),
]

for entity_name, attrs in entities:
    heading(entity_name, level=3)
    add_table(
        ["Attribute", "Type", "Constraint"],
        attrs
    )
    doc.add_paragraph()  # spacer

heading("E-R Relationship Diagram", level=3)
code_block(
    "+----------+       teaches        +----------+      belongs to    +----------+\n"
    "|   User   |----(1)---------(M)--|  Course   |----(M)--------(1)--|Semester  |\n"
    "| (PK: id) |                     | (PK: id) |                    | (PK: id) |\n"
    "+----------+                     +----------+                    +----------+\n"
    "  |  |  |  |                      |  |  |                             |\n"
    "  |  |  |  +-(M)--[Enrollment]--(M)+  |  |                             |\n"
    "  |  |  |                             |  |                             |\n"
    "  |  |  +---(M)--[Waitlist]-----(M)---+  |                             |\n"
    "  |  |                                   |                             |\n"
    "  |  +-----(M)--[Review]----------(M)----+                             |\n"
    "  |                                                                    |\n"
    "  +-------(M)--[Complaint]--(M)---User (target)                       |\n"
    "  +-------(1)--[Warning]--(M)                                         |\n"
    "  +-------(1)--[Application]--(M)                                     |\n"
    "  +-------(1)--[GraduationRequest]--(M)                               |\n"
    "  +-------(M)--[HonorRoll]--(M)---Semester                            |\n"
    "                                                                       \n"
    "               [TabooWord] (standalone, no FK)                         "
)

heading("Relationship Summary", level=3)
relationships = [
    "User (1) --- (M) Course: instructor teaches courses",
    "User (M) --- (M) Course via Enrollment: students enroll in courses",
    "User (M) --- (M) Course via Waitlist: students waitlisted for courses",
    "User (M) --- (M) Course via Review: students review courses",
    "User (1) --- (M) Warning: user receives warnings",
    "User (1) --- (M) Application: user submits applications",
    "User (1) --- (M) GraduationRequest: student applies for graduation",
    "User (M) --- (M) User via Complaint: filer complains about target",
    "User (M) --- (M) Semester via HonorRoll: student honored per semester",
    "Semester (1) --- (M) Course: semester contains courses",
    "TabooWord: standalone entity with no foreign keys",
]
for r in relationships:
    bullet(r)

doc.add_page_break()

# ════════════════════════════════════════
# 4. DETAILED DESIGN (PSEUDO-CODE)
# ════════════════════════════════════════
heading("4. Detailed Design", level=1)
para("This section provides pseudo-code for every method in the system, organized by subsystem. "
     "Each method specifies its input, output, and main logic.")

pseudocode_sections = [
    ("4.1 Session Management", [
        ("createSession(userId)",
         "FUNCTION createSession(userId: Integer) -> SessionData | null\n"
         "  INPUT: userId -- the database ID of the authenticated user\n"
         "  OUTPUT: SessionData object or null if user not found\n"
         "\n"
         "  user = Database.User.findUnique(WHERE id = userId)\n"
         "  IF user IS null THEN\n"
         "    RETURN null\n"
         "  END IF\n"
         "\n"
         "  session = {\n"
         "    userId: user.id,\n"
         "    role: user.role,\n"
         "    firstName: user.firstName,\n"
         "    lastName: user.lastName,\n"
         "    email: user.email\n"
         "  }\n"
         "\n"
         "  SET httpOnly cookie \"session\" = JSON.stringify(session)\n"
         "    WITH maxAge = 86400 seconds (1 day)\n"
         "    WITH sameSite = \"lax\", path = \"/\"\n"
         "\n"
         "  RETURN session\n"
         "END FUNCTION"),
        ("getSession()",
         "FUNCTION getSession() -> SessionData | null\n"
         "  INPUT: none (reads from HTTP cookies)\n"
         "  OUTPUT: SessionData object or null\n"
         "\n"
         "  raw = READ cookie \"session\"\n"
         "  IF raw IS empty THEN RETURN null\n"
         "\n"
         "  TRY\n"
         "    RETURN JSON.parse(raw) AS SessionData\n"
         "  CATCH\n"
         "    RETURN null\n"
         "  END TRY\n"
         "END FUNCTION"),
        ("destroySession()",
         "FUNCTION destroySession() -> void\n"
         "  INPUT: none\n"
         "  OUTPUT: none (deletes session cookie)\n"
         "\n"
         "  DELETE cookie \"session\"\n"
         "END FUNCTION"),
    ]),
    ("4.2 Authentication", [
        ("loginUser(email, password)",
         "FUNCTION loginUser(email: String, password: String) -> LoginResult\n"
         "  INPUT: email, password\n"
         "  OUTPUT: { success, error?, role?, firstName?, redirect? }\n"
         "\n"
         "  user = Database.User.findUnique(WHERE email = email.trim().toLowerCase())\n"
         "\n"
         "  IF user IS null OR user.password != password THEN\n"
         "    RETURN { success: false, error: \"Invalid email or password.\" }\n"
         "  END IF\n"
         "  IF user.terminated THEN RETURN { success: false, error: \"Account terminated.\" }\n"
         "  IF user.suspended THEN RETURN { success: false, error: \"Account suspended.\" }\n"
         "  IF user.fired THEN RETURN { success: false, error: \"Account deactivated.\" }\n"
         "\n"
         "  CALL createSession(user.id)\n"
         "  redirect = user.mustChangePassword ? \"/change-password\" : \"/dashboard\"\n"
         "  RETURN { success: true, role: user.role, firstName: user.firstName, redirect }\n"
         "END FUNCTION"),
        ("logoutUser()",
         "FUNCTION logoutUser() -> void\n"
         "  CALL destroySession()\n"
         "END FUNCTION"),
        ("changePassword(request)",
         "FUNCTION changePassword(request: HTTP POST) -> JSON\n"
         "  INPUT: { newPassword }\n"
         "  OUTPUT: { success, error?, redirect? }\n"
         "\n"
         "  session = CALL getSession()\n"
         "  IF session IS null THEN RETURN HTTP 401\n"
         "\n"
         "  IF newPassword IS empty OR length < 6 THEN\n"
         "    RETURN { success: false, error: \"Password must be >= 6 chars.\" }\n"
         "  END IF\n"
         "\n"
         "  Database.User.update(WHERE id = session.userId,\n"
         "    SET password = newPassword, mustChangePassword = false)\n"
         "  CALL createSession(session.userId)  // refresh\n"
         "  RETURN { success: true, redirect: \"/dashboard\" }\n"
         "END FUNCTION"),
    ]),
    ("4.3 Dashboard Routing", [
        ("Dashboard() -- main router",
         "FUNCTION Dashboard() -> HTML Page\n"
         "  session = CALL getSession()\n"
         "  IF session IS null THEN REDIRECT to \"/login\"\n"
         "\n"
         "  SWITCH session.role\n"
         "    CASE \"REGISTRAR\": cards = registrarCards (8 cards)\n"
         "    CASE \"INSTRUCTOR\": cards = instructorCards (5 cards)\n"
         "    DEFAULT: cards = studentCards (6 cards)\n"
         "  END SWITCH\n"
         "\n"
         "  RENDER page with nav bar, welcome, card grid\n"
         "END FUNCTION"),
        ("StudentDashboard()",
         "FUNCTION StudentDashboard() -> HTML Page\n"
         "  session = CALL getSession()\n"
         "  IF session IS null OR role != \"STUDENT\" THEN REDIRECT \"/login\"\n"
         "\n"
         "  user = Database.User.findUnique(WHERE id = session.userId,\n"
         "    INCLUDE enrollments->course->instructor,semester,\n"
         "    INCLUDE warningsReceived(removed=false), honorRollEntries)\n"
         "\n"
         "  currentSemester = Database.Semester.findFirst(WHERE isCurrent=true)\n"
         "  currentEnrollments = FILTER enrollments WHERE semester.isCurrent\n"
         "  pastEnrollments = FILTER enrollments WHERE NOT semester.isCurrent\n"
         "\n"
         "  RENDER: nav, welcome banner, stats (GPA, courses, warnings, fines),\n"
         "    honor roll badge, current courses, past courses, sidebar\n"
         "END FUNCTION"),
        ("InstructorDashboard()",
         "FUNCTION InstructorDashboard() -> HTML Page\n"
         "  session = CALL getSession()\n"
         "  IF session IS null OR role != \"INSTRUCTOR\" THEN REDIRECT \"/login\"\n"
         "\n"
         "  user = Database.User.findUnique(WHERE id = session.userId,\n"
         "    INCLUDE warningsReceived(removed=false))\n"
         "  courses = Database.Course.findMany(\n"
         "    WHERE instructorId=user.id AND semester.isCurrent,\n"
         "    INCLUDE enrollments->user, reviews, semester)\n"
         "  totalStudents = SUM(enrollments.length) across courses\n"
         "\n"
         "  FOR EACH course: avgRating = mean(reviews.rating) or null\n"
         "\n"
         "  RENDER: nav, welcome, stats, course cards with student lists, warnings\n"
         "END FUNCTION"),
        ("RegistrarDashboard()",
         "FUNCTION RegistrarDashboard() -> HTML Page\n"
         "  session = CALL getSession()\n"
         "  IF session IS null OR role != \"REGISTRAR\" THEN REDIRECT \"/login\"\n"
         "\n"
         "  PARALLEL FETCH:\n"
         "    totalStudents, totalInstructors, totalCourses,\n"
         "    pendingApplications, pendingComplaints, pendingGraduations,\n"
         "    currentSemester, activeWarnings, suspendedStudents\n"
         "\n"
         "  topStudents = Database.User.findMany(\n"
         "    WHERE role=STUDENT AND NOT terminated, ORDER BY gpa DESC, LIMIT 5)\n"
         "\n"
         "  RENDER: nav, welcome, stats, pending items, management grid,\n"
         "    top students sidebar, suspended alert\n"
         "END FUNCTION"),
    ]),
    ("4.4 Course Registration", [
        ("registerForCourse(studentId, courseId)",
         "FUNCTION registerForCourse(studentId, courseId) -> Result\n"
         "  INPUT: studentId, courseId\n"
         "  OUTPUT: { success, message, waitlisted? }\n"
         "\n"
         "  student = Database.User.findUnique(WHERE id = studentId)\n"
         "  IF student.suspended OR student.terminated THEN\n"
         "    RETURN { success: false, message: \"Account restricted.\" }\n"
         "\n"
         "  course = Database.Course.findUnique(WHERE id=courseId, INCLUDE enrollments,semester)\n"
         "  IF course.semester.period != \"REGISTRATION\" THEN\n"
         "    RETURN { success: false, message: \"Registration not open.\" }\n"
         "\n"
         "  currentCount = COUNT enrollments WHERE userId=studentId AND semester.isCurrent\n"
         "  IF currentCount >= 4 THEN RETURN error \"Maximum 4 courses\"\n"
         "\n"
         "  // Check time conflict\n"
         "  FOR EACH existing enrollment:\n"
         "    IF schedule OVERLAPS THEN RETURN error \"Time conflict\"\n"
         "\n"
         "  // Check retake eligibility\n"
         "  IF previous enrollment exists with non-F grade THEN\n"
         "    RETURN error \"Cannot retake unless previously failed with F\"\n"
         "\n"
         "  // Check capacity\n"
         "  IF enrollments.length >= maxStudents THEN\n"
         "    CREATE Waitlist entry with next position\n"
         "    RETURN { success: true, waitlisted: true }\n"
         "\n"
         "  CREATE Enrollment(userId, courseId, status=ENROLLED)\n"
         "  RETURN { success: true, message: \"Enrolled.\" }\n"
         "END FUNCTION"),
    ]),
    ("4.5 Grade Assignment", [
        ("assignGrades(courseId, grades)",
         "FUNCTION assignGrades(courseId, grades[]) -> Result\n"
         "  INPUT: courseId, array of { studentId, grade }\n"
         "  OUTPUT: { success, message }\n"
         "\n"
         "  session = CALL getSession()\n"
         "  IF role != INSTRUCTOR THEN RETURN unauthorized\n"
         "\n"
         "  course = Database.Course.findUnique(\n"
         "    WHERE id=courseId AND instructorId=session.userId)\n"
         "  IF course.semester.period != \"GRADING\" THEN RETURN error\n"
         "\n"
         "  FOR EACH { studentId, grade } IN grades:\n"
         "    IF grade NOT IN [A,B,C,D,F] THEN RETURN error\n"
         "    UPDATE Enrollment SET grade=grade, status=COMPLETED\n"
         "      WHERE userId=studentId AND courseId=courseId\n"
         "\n"
         "  RETURN { success: true }\n"
         "END FUNCTION"),
    ]),
    ("4.6 Academic Standing Evaluation", [
        ("evaluateAcademicStanding(semesterId)",
         "FUNCTION evaluateAcademicStanding(semesterId) -> void\n"
         "  gradePoints = { A:4.0, B:3.0, C:2.0, D:1.0, F:0.0 }\n"
         "\n"
         "  FOR EACH active student:\n"
         "    Calculate semesterGPA and cumulativeGPA\n"
         "    UPDATE user.gpa = cumulativeGPA\n"
         "\n"
         "    // Termination checks\n"
         "    IF any course failed twice (same code) THEN SET terminated=true; CONTINUE\n"
         "    IF cumulativeGPA < 2.0 THEN SET terminated=true; CONTINUE\n"
         "\n"
         "    // Probation warning\n"
         "    IF cumulativeGPA BETWEEN 2.0 AND 2.25 THEN\n"
         "      CREATE Warning(\"Low GPA probation\")\n"
         "\n"
         "    // Honor roll\n"
         "    IF semesterGPA > 3.75 THEN CREATE HonorRoll(type=SEMESTER)\n"
         "    IF cumulativeGPA > 3.5 AND completedSemesters > 1 THEN\n"
         "      CREATE HonorRoll(type=OVERALL)\n"
         "\n"
         "    // Honor removes warning\n"
         "    FOR EACH unused honor:\n"
         "      IF active warning exists THEN\n"
         "        SET warning.removed=true, honor.usedToRemoveWarning=true\n"
         "        DECREMENT user.warnings\n"
         "\n"
         "    // Suspension check\n"
         "    IF warnings >= 3 THEN SET suspended=true, fineOwed += 100\n"
         "END FUNCTION"),
    ]),
    ("4.7 Course Review Submission", [
        ("submitReview(authorId, courseId, rating, comment)",
         "FUNCTION submitReview(authorId, courseId, rating, comment) -> Result\n"
         "  Verify enrollment exists AND grade IS null\n"
         "  Check no duplicate review (unique authorId+courseId)\n"
         "\n"
         "  // Taboo word filtering\n"
         "  tabooWords = Database.TabooWord.findMany()\n"
         "  tabooCount = 0\n"
         "  FOR EACH taboo IN tabooWords:\n"
         "    IF comment CONTAINS taboo.word THEN\n"
         "      tabooCount++; REPLACE word WITH \"***\"\n"
         "\n"
         "  IF tabooCount >= 3: hidden=true, warnings=2\n"
         "  ELSE IF tabooCount >= 1: warnings=1\n"
         "  ELSE: warnings=0\n"
         "\n"
         "  CREATE Review(authorId, courseId, rating, comment, hidden)\n"
         "  IF warnings > 0: CREATE Warning records, INCREMENT user.warnings\n"
         "\n"
         "  // Check course avg rating\n"
         "  avgRating = mean of visible reviews\n"
         "  IF avgRating < 2.0 THEN CREATE Warning for instructor\n"
         "\n"
         "  RETURN { success: true, warningsIssued }\n"
         "END FUNCTION"),
    ]),
    ("4.8 Complaint Management", [
        ("fileComplaint(filerId, targetId, description)",
         "FUNCTION fileComplaint(filerId, targetId, description) -> Result\n"
         "  IF filerId = targetId THEN RETURN error\n"
         "  IF description IS empty THEN RETURN error\n"
         "\n"
         "  CREATE Complaint(filerId, targetId, description, status=PENDING)\n"
         "  RETURN { success: true }\n"
         "END FUNCTION"),
        ("processComplaint(complaintId, action, resolution)",
         "FUNCTION processComplaint(complaintId, action, resolution) -> Result\n"
         "  complaint = findUnique(complaintId, INCLUDE filer, target)\n"
         "\n"
         "  IF action = DISMISS:\n"
         "    UPDATE status=DISMISSED; RETURN\n"
         "\n"
         "  IF action = WARN:\n"
         "    CREATE Warning for target; INCREMENT target.warnings\n"
         "\n"
         "  IF action = DEREGISTER:\n"
         "    DELETE target's current enrollments\n"
         "    CREATE Warning for target\n"
         "\n"
         "  IF action = WARN_FILER:\n"
         "    CREATE Warning for filer (unjustified complaint)\n"
         "\n"
         "  UPDATE complaint status=RESOLVED\n"
         "\n"
         "  // Check 3-warning suspension\n"
         "  IF target.warnings >= 3 AND NOT suspended THEN\n"
         "    SET suspended=true, fineOwed += 100\n"
         "END FUNCTION"),
    ]),
    ("4.9 Application Management", [
        ("submitApplication(userId, type, priorGpa, justification)",
         "FUNCTION submitApplication(userId, type, priorGpa, justification) -> Result\n"
         "  IF type=STUDENT AND (priorGpa null OR justification empty) THEN\n"
         "    RETURN error \"Student apps require GPA and justification\"\n"
         "  IF pending application exists THEN RETURN error \"Already pending\"\n"
         "\n"
         "  CREATE Application(userId, type, status=PENDING, priorGpa, justification)\n"
         "  RETURN { success: true }\n"
         "END FUNCTION"),
        ("reviewApplication(appId, decision, justification)",
         "FUNCTION reviewApplication(appId, decision, justification) -> Result\n"
         "  app = findUnique(appId, INCLUDE user)\n"
         "\n"
         "  IF decision = REJECT:\n"
         "    IF app.type=STUDENT AND priorGpa>3.0 AND no justification THEN\n"
         "      RETURN error \"Must justify rejecting qualified student\"\n"
         "    UPDATE status=REJECTED; RETURN\n"
         "\n"
         "  // ACCEPT\n"
         "  IF type = STUDENT:\n"
         "    email = firstName[0] + lastName + \"00@cuny.edu\"\n"
         "    tempPassword = generateRandom()\n"
         "    UPDATE User SET role=STUDENT, email, password=temp, mustChangePassword=true\n"
         "  ELSE:\n"
         "    UPDATE User SET role=INSTRUCTOR\n"
         "\n"
         "  UPDATE Application status=ACCEPTED\n"
         "END FUNCTION"),
    ]),
    ("4.10 Graduation", [
        ("applyForGraduation(studentId)",
         "FUNCTION applyForGraduation(studentId) -> Result\n"
         "  completedCourses = COUNT Enrollment WHERE userId=studentId\n"
         "    AND grade NOT null AND grade != F\n"
         "  isReckless = completedCourses < 8\n"
         "\n"
         "  CREATE GraduationRequest(userId=studentId, status=PENDING)\n"
         "\n"
         "  IF isReckless THEN\n"
         "    RETURN { success: true, warning: \"Fewer than 8 courses\" }\n"
         "  RETURN { success: true }\n"
         "END FUNCTION"),
        ("reviewGraduation(requestId, decision)",
         "FUNCTION reviewGraduation(requestId, decision) -> Result\n"
         "  request = findUnique(requestId, INCLUDE user)\n"
         "  completedPassing = COUNT Enrollment WHERE grade IN [A,B,C,D]\n"
         "\n"
         "  IF decision = APPROVE:\n"
         "    IF completedPassing < 8 OR suspended OR fineOwed > 0 THEN\n"
         "      RETURN error \"Does not meet requirements\"\n"
         "    SET graduated=true; UPDATE status=APPROVED\n"
         "\n"
         "  IF decision = REJECT:\n"
         "    UPDATE status=REJECTED\n"
         "    IF completedPassing < 8 THEN\n"
         "      CREATE Warning(\"Reckless graduation application\")\n"
         "END FUNCTION"),
    ]),
    ("4.11 Semester and Running Rules", [
        ("advanceSemesterPeriod(semesterId)",
         "FUNCTION advanceSemesterPeriod(semesterId) -> Result\n"
         "  order = [CLASS_SETUP, REGISTRATION, RUNNING, GRADING, COMPLETED]\n"
         "  currentIndex = order.indexOf(semester.period)\n"
         "  IF currentIndex = last THEN RETURN error \"Already completed\"\n"
         "\n"
         "  newPeriod = order[currentIndex + 1]\n"
         "  UPDATE semester SET period = newPeriod\n"
         "\n"
         "  IF newPeriod = RUNNING THEN CALL enforceRunningPeriodRules()\n"
         "  IF newPeriod = COMPLETED THEN CALL evaluateAcademicStanding()\n"
         "END FUNCTION"),
        ("enforceRunningPeriodRules(semesterId)",
         "FUNCTION enforceRunningPeriodRules(semesterId) -> void\n"
         "  // 1. Warn students with < 2 courses\n"
         "  FOR EACH active student:\n"
         "    IF enrollmentCount > 0 AND enrollmentCount < 2 THEN\n"
         "      CREATE Warning; INCREMENT warnings\n"
         "\n"
         "  // 2. Cancel courses with < 3 students\n"
         "  FOR EACH non-cancelled course:\n"
         "    IF enrollments.length < 3 THEN\n"
         "      SET cancelled=true\n"
         "      CREATE Warning for instructor\n"
         "      DELETE enrollments; flag students for special registration\n"
         "\n"
         "  // 3. Suspend instructors who lost ALL courses\n"
         "  FOR EACH instructor:\n"
         "    IF had courses but all cancelled THEN\n"
         "      SET suspended=true (cannot teach next semester)\n"
         "END FUNCTION"),
    ]),
    ("4.12 Taboo Word Management", [
        ("addTabooWord(word)",
         "FUNCTION addTabooWord(word) -> Result\n"
         "  IF word IS empty THEN RETURN error\n"
         "  IF word already exists (unique constraint) THEN RETURN error\n"
         "  CREATE TabooWord(word = word.toLowerCase())\n"
         "  RETURN { success: true }\n"
         "END FUNCTION"),
        ("removeTabooWord(wordId)",
         "FUNCTION removeTabooWord(wordId) -> Result\n"
         "  DELETE TabooWord WHERE id = wordId\n"
         "  RETURN { success: true }\n"
         "END FUNCTION"),
    ]),
    ("4.13 Waitlist Management", [
        ("admitFromWaitlist(courseId, studentId)",
         "FUNCTION admitFromWaitlist(courseId, studentId) -> Result\n"
         "  entry = Waitlist.findUnique(userId+courseId)\n"
         "  IF not found OR status != WAITING THEN RETURN error\n"
         "\n"
         "  student = User.findUnique(studentId)\n"
         "  IF suspended OR terminated THEN RETURN error \"No longer eligible\"\n"
         "\n"
         "  CREATE Enrollment(userId, courseId, status=ENROLLED)\n"
         "  UPDATE Waitlist SET status=ADMITTED\n"
         "\n"
         "  // Reorder remaining positions\n"
         "  remaining = Waitlist.findMany(courseId, status=WAITING, ORDER BY position)\n"
         "  FOR i = 0 TO remaining.length-1:\n"
         "    UPDATE remaining[i].position = i + 1\n"
         "END FUNCTION"),
    ]),
    ("4.14 Course Setup", [
        ("createCourse(...)",
         "FUNCTION createCourse(code, name, credits, schedule, maxStudents,\n"
         "                      instructorId, semesterId) -> Result\n"
         "  IF semester.period != CLASS_SETUP THEN RETURN error\n"
         "  IF course code exists in same semester THEN RETURN error \"Duplicate\"\n"
         "\n"
         "  CREATE Course(code, name, credits, schedule, maxStudents,\n"
         "                instructorId, semesterId)\n"
         "  RETURN { success: true, course }\n"
         "END FUNCTION"),
        ("cancelCourse(courseId)",
         "FUNCTION cancelCourse(courseId) -> Result\n"
         "  UPDATE Course SET cancelled=true\n"
         "  IF instructorId exists THEN\n"
         "    CREATE Warning for instructor\n"
         "    INCREMENT instructor.warnings\n"
         "  RETURN { success: true }\n"
         "END FUNCTION"),
    ]),
    ("4.15 AI Assistant (Lite)", [
        ("askAI(question)",
         "FUNCTION askAI(question) -> Result\n"
         "  IF question IS empty THEN RETURN \"Please enter a question.\"\n"
         "\n"
         "  // RAG: search local knowledge base\n"
         "  policies = LOAD \"data/policy.json\"\n"
         "  keywords = TOKENIZE(question.toLowerCase())\n"
         "  matchedPolicies = []\n"
         "  FOR EACH policy IN policies:\n"
         "    score = COUNT matching keywords\n"
         "    IF score > 0 THEN matchedPolicies.push({ policy, score })\n"
         "  SORT by score DESC\n"
         "\n"
         "  // Build prompt\n"
         "  IF matches found THEN\n"
         "    context = TOP 3 policies joined as text\n"
         "    prompt = \"Based on these policies:\\n\" + context + \"\\nAnswer: \" + question\n"
         "    source = \"knowledge_base\"\n"
         "  ELSE\n"
         "    prompt = question\n"
         "    source = \"general_llm\"\n"
         "\n"
         "  // Call OpenAI\n"
         "  TRY\n"
         "    response = OpenAI.chat(model=\"gpt-4o-mini\", prompt)\n"
         "    answer = response.choices[0].message.content\n"
         "  CATCH\n"
         "    answer = \"AI assistant is currently unavailable.\"\n"
         "    source = \"error\"\n"
         "\n"
         "  RETURN { answer, source }\n"
         "END FUNCTION"),
    ]),
]

for section_title, methods in pseudocode_sections:
    heading(section_title, level=2)
    for method_name, pseudo in methods:
        heading(method_name, level=3)
        code_block(pseudo)

doc.add_page_break()

# ════════════════════════════════════════
# 5. SYSTEM SCREENS
# ════════════════════════════════════════
heading("5. System Screens", level=1)

screens = [
    ("5.1 Login Page",
     "The login page presents a centered form with the CUNYZeroLite branding in navy blue. "
     "Fields include email and password inputs with a \"Sign In\" button. Error messages appear inline "
     "(e.g., \"Invalid email or password,\" \"Your account has been terminated\"). The page features "
     "the midnight navy gradient background consistent with the university portal theme."),
    ("5.2 Change Password Page",
     "Displayed when mustChangePassword = true after first login. Shows a simple form with "
     "\"New Password\" and \"Confirm Password\" fields. Validates minimum 6 characters. "
     "On success, redirects to the role-based dashboard."),
    ("5.3 Main Dashboard (Card Grid)",
     "After login, users see a card grid layout tailored to their role:\n"
     "- Students: 6 cards (Profile, Course Registration, Schedule Builder, Grades, Reviews, Complaints)\n"
     "- Instructors: 5 cards (Profile, My Courses, Grade Students, My Students, File Complaint)\n"
     "- Registrar: 8 cards (Profile, Semester Mgmt, Course Mgmt, Students, Apps, Complaints, Graduations, Taboo Words)\n"
     "Each card shows an icon, title, and subtitle. Uses 3-column grid on desktop, responsive on mobile."),
    ("5.4 Student Dashboard (Sample Prototype)",
     "The student dashboard is the most comprehensive screen and serves as the sample prototype. It features:\n\n"
     "- Navigation bar: Dark navy (#0f172a) with CunyZeroLite logo, \"Student Portal\" subtitle, student name, logout button.\n"
     "- Welcome banner: Gradient navy panel showing student's first name and current semester period.\n"
     "- Stats row: Four white cards displaying GPA, current course count, warnings count, and fine balance.\n"
     "- Honor roll badge: Golden amber banner shown only for honor roll students.\n"
     "- Current courses: Cards with course code, name, schedule, credits, instructor, enrollment status.\n"
     "- Past courses: Compact list with grades.\n"
     "- Sidebar: Quick action links (Register, Reviews, Complaint, Graduation), active warnings list, profile card."),
    ("5.5 Instructor Dashboard",
     "Emerald-themed navigation with \"Instructor Portal\" subtitle. Three stat cards (My Courses, "
     "Total Students, Warnings). Each course displayed with dark header showing code, name, schedule, "
     "average review rating, and active/cancelled badge. Enrolled student list shows name, GPA, email, "
     "and current grade status."),
    ("5.6 Registrar Dashboard",
     "Red-themed navigation with shield icon and \"Registrar Portal\" subtitle. Four stat cards "
     "(Students, Instructors, Active Courses, Active Warnings). Three clickable pending-item cards "
     "(Applications, Complaints, Graduations). Management action grid with four cards linking to "
     "Semester, Course, Student, and Taboo Word management. Top 5 students sidebar ranked by GPA. "
     "Suspended students alert banner."),
]

for title, description in screens:
    heading(title, level=2)
    para(description)

doc.add_page_break()

# ════════════════════════════════════════
# 6. MEETING MEMOS
# ════════════════════════════════════════
heading("6. Meeting Memos", level=1)

meetings = [
    ("Meeting 1 \u2014 March 15, 2026",
     "Attendees: Diego Reyes Liranzo, Daniel Olekszyk, Samia Islam, Maisha Islam, Kyle Gosine",
     [
         "Reviewed project specification and identified all 17 use cases.",
         "Decided on technology stack: Next.js 16, TypeScript, Prisma 7.5, SQLite, Tailwind CSS 4.",
         "Assigned initial responsibilities: Diego on login page, Samia on database schema, Daniel on API routes, Maisha on UI components, Kyle on documentation.",
         "Agreed on Git workflow: feature branches, PRs to main, code review required.",
     ],
     [
         "SQLite limitations for concurrent access; decided acceptable for demo scope.",
         "Need to clarify taboo word filtering edge cases with professor.",
     ]),
    ("Meeting 2 \u2014 March 24, 2026",
     "Attendees: Diego Reyes Liranzo, Daniel Olekszyk, Samia Islam, Maisha Islam, Kyle Gosine",
     [
         "Completed Phase I SRS document (Version 1.0).",
         "Finalized all use-case descriptions and supplementary requirements.",
         "Reviewed database schema design (11 models).",
         "Diego demonstrated working login page prototype.",
     ],
     [
         "OpenAI API key management \u2014 agreed to use .env file, not commit keys.",
         "Timeline tight for Phase II; began planning pseudo-code assignments.",
     ]),
    ("Meeting 3 \u2014 April 5, 2026",
     "Attendees: Diego Reyes Liranzo, Daniel Olekszyk, Samia Islam, Maisha Islam, Kyle Gosine",
     [
         "Database schema merged to main (11 models, seed script with 16 users).",
         "Samia implemented authentication system, session management, and role-based dashboards.",
         "Integrated portal logic with navy UI theme.",
         "Reviewed student, instructor, and registrar dashboard prototypes.",
     ],
     [
         "Some features (course registration, grading, reviews) still need implementation.",
         "Need to coordinate on complaint processing and graduation workflows.",
     ]),
    ("Meeting 4 \u2014 April 18, 2026",
     "Attendees: Diego Reyes Liranzo, Daniel Olekszyk, Samia Islam, Maisha Islam, Kyle Gosine",
     [
         "Began drafting Phase II Design Report.",
         "Assigned diagram creation: Samia on E-R diagram and pseudo-code, Diego on sequence diagrams, Daniel on Petri-nets.",
         "Reviewed all use-case scenarios for completeness.",
         "Planned remaining implementation for Phase III.",
     ],
     [
         "Phase II deadline April 23; need to finalize report in 2 days.",
         "AI chat feature depends on valid OpenAI API key \u2014 need backup plan.",
     ]),
]

for title, attendees, topics, concerns in meetings:
    heading(title, level=2)
    para(attendees, italic=True)
    para("Topics:", bold=True)
    for t in topics:
        bullet(t)
    para("Concerns:", bold=True)
    for c in concerns:
        bullet(c)

doc.add_page_break()

# ════════════════════════════════════════
# 7. REPOSITORY
# ════════════════════════════════════════
heading("7. Repository", level=1)

para("GitHub Repository: https://github.com/dreyesliranzo/CunyZeroLite", bold=True)

para("Branch structure:")
bullets_repo = [
    "main \u2014 stable, reviewed code only",
    "feature/database-schema \u2014 Prisma schema and seed data (merged)",
    "feature/auth-dashboards \u2014 Authentication, session management, role-based dashboards",
    "feature/student-portal \u2014 Student portal UI (merged)",
    "feature/ai-chat \u2014 AI assistant integration (in progress)",
]
for b in bullets_repo:
    bullet(b)

para("Technology Stack:")
tech = [
    "Next.js 16 (React framework)",
    "TypeScript (type safety)",
    "Prisma 7.5 (ORM)",
    "SQLite via better-sqlite3 (database)",
    "Tailwind CSS 4 (styling)",
    "OpenAI GPT-4o-mini (AI assistant)",
]
for t in tech:
    bullet(t)

# ── Header/Footer ──
add_header_footer()

# ── SAVE ──
output_path = "/home/samia/CunyZeroLite/docs/GroupK_CunyZero_PhaseII.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
