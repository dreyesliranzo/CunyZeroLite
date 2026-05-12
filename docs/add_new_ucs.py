#!/usr/bin/env python3
"""Add UC-18, UC-19, UC-20 to the Phase II report and update counts/references."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import copy, re

doc = Document("/home/samia/CunyZeroLite/docs/GroupK_CunyZero_PhaseII.docx")

style = doc.styles['Normal']

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

# ── Fix counts in existing text: find "seventeen" and "17" references ──
# We'll update these by adding content at the end before Meeting Memos

# First, let's update the revision history and intro text by finding and replacing in paragraphs
for p in doc.paragraphs:
    for run in p.runs:
        if "seventeen use cases" in run.text:
            run.text = run.text.replace("seventeen use cases", "twenty use cases")
        if "all seventeen use cases" in run.text:
            run.text = run.text.replace("all seventeen use cases", "all twenty use cases")
        if "17 use cases" in run.text.lower():
            run.text = run.text.replace("17", "20")

# ════════════════════════════════════════
# Now add the new use cases, diagrams, pseudo-code
# We'll add them at the end of the document (before we re-save)
# ════════════════════════════════════════

doc.add_page_break()

# ════════════════════════════════════════
# UC-18: Drop / Withdraw from Course
# ════════════════════════════════════════
heading("Addendum: Additional Use Cases (UC-18 through UC-20)", level=1)
para("The following three use cases extend the original seventeen to cover course drop/withdrawal, "
     "fine payment, and schedule planning functionality. These are derived from the system's existing "
     "data model (Enrollment.status, User.fineOwed, Course.schedule) and planned feature roadmap.")

doc.add_paragraph()

heading("UC-18: Drop / Withdraw from Course", level=2)

heading("Use-Case Scenario", level=3)
para("Normal Scenario \u2014 Drop (within first 2 weeks):", bold=True)
bullet("1. Student clicks \"Drop Course\" within the drop window.")
bullet("2. System verifies current date is within the drop deadline (first 2 weeks of RUNNING period).")
bullet("3. System deletes the Enrollment record entirely (no transcript entry).")
bullet("4. If students are on the waitlist, the first waitlisted student is auto-promoted to ENROLLED.")
bullet("5. System returns success; course removed from student's schedule.")

para("Normal Scenario \u2014 Withdraw (after drop deadline):", bold=True)
bullet("1. Student clicks \"Withdraw from Course\" after the drop deadline has passed.")
bullet("2. System verifies semester is in RUNNING period (not COMPLETED).")
bullet("3. System updates Enrollment: status = \"WITHDRAWN\", grade = \"W\".")
bullet("4. W appears on transcript but does not affect GPA.")
bullet("5. System checks if student drops below minimum course load (2 courses) and warns.")

para("Exceptional Scenarios:", bold=True)
bullet("Drop window closed: Date past 2-week deadline \u2192 system redirects to Withdraw flow instead.")
bullet("Semester COMPLETED: Cannot drop or withdraw after semester ends \u2192 400 error.")
bullet("Not enrolled: Enrollment record not found \u2192 404 error.")
bullet("Below minimum load: Dropping below 2 courses triggers a warning to the student.")

heading("Sequence Diagram \u2014 UC-18", level=3)
code_block(
    "Student          Portal            API               Database\n"
    "  |                |                 |                  |\n"
    "  |-- Drop Course->|                 |                  |\n"
    "  |                |-- DELETE        |                  |\n"
    "  |                |   /enroll ----->|                  |\n"
    "  |                |                 |-- check drop     |\n"
    "  |                |                 |   window         |\n"
    "  |                |                 |                  |\n"
    "  |          [IF within drop window]                    |\n"
    "  |                |                 |-- DELETE         |\n"
    "  |                |                 |   Enrollment --->|\n"
    "  |                |                 |-- promote        |\n"
    "  |                |                 |   waitlist[0] -->|\n"
    "  |                |                 |<-- OK -----------|\n"
    "  |                |<-- 200 dropped -|                  |\n"
    "  |<-- removed ----|                 |                  |\n"
    "  |                                                     |\n"
    "  |          [IF after drop deadline]                   |\n"
    "  |                |                 |-- UPDATE         |\n"
    "  |                |                 |   status=WITH-   |\n"
    "  |                |                 |   DRAWN, grade=W>|\n"
    "  |                |                 |<-- OK -----------|\n"
    "  |                |<-- 200 W grade -|                  |\n"
    "  |<-- W on record |                 |                  |"
)

heading("Pseudo-Code \u2014 dropOrWithdrawCourse(studentId, courseId)", level=3)
code_block(
    "FUNCTION dropOrWithdrawCourse(studentId, courseId) -> Result\n"
    "  INPUT: studentId, courseId\n"
    "  OUTPUT: { success, action, message }\n"
    "\n"
    "  enrollment = Database.Enrollment.findUnique(\n"
    "    WHERE userId = studentId AND courseId = courseId)\n"
    "  IF enrollment IS null THEN\n"
    "    RETURN { success: false, message: \"Not enrolled.\" }\n"
    "\n"
    "  semester = Database.Semester.findFirst(WHERE isCurrent = true)\n"
    "  IF semester.period = \"COMPLETED\" THEN\n"
    "    RETURN { success: false, message: \"Cannot modify after semester end.\" }\n"
    "\n"
    "  // Determine if within drop window (first 2 weeks of RUNNING)\n"
    "  dropDeadline = semester.startDate + 14 days\n"
    "  currentDate = NOW()\n"
    "\n"
    "  IF currentDate <= dropDeadline AND semester.period = \"RUNNING\" THEN\n"
    "    // DROP: delete enrollment entirely\n"
    "    Database.Enrollment.delete(WHERE id = enrollment.id)\n"
    "\n"
    "    // Promote first waitlisted student\n"
    "    nextWaitlist = Database.Waitlist.findFirst(\n"
    "      WHERE courseId = courseId AND status = \"WAITING\",\n"
    "      ORDER BY position ASC)\n"
    "    IF nextWaitlist EXISTS THEN\n"
    "      Database.Enrollment.create(\n"
    "        userId = nextWaitlist.userId, courseId, status = \"ENROLLED\")\n"
    "      Database.Waitlist.update(\n"
    "        WHERE id = nextWaitlist.id, SET status = \"ADMITTED\")\n"
    "    END IF\n"
    "\n"
    "    RETURN { success: true, action: \"DROPPED\", message: \"Course dropped.\" }\n"
    "\n"
    "  ELSE IF semester.period = \"RUNNING\" THEN\n"
    "    // WITHDRAW: keep record with W grade\n"
    "    Database.Enrollment.update(\n"
    "      WHERE id = enrollment.id,\n"
    "      SET status = \"WITHDRAWN\", grade = \"W\")\n"
    "\n"
    "    // Check minimum course load\n"
    "    remainingCount = Database.Enrollment.count(\n"
    "      WHERE userId = studentId AND course.semester.isCurrent\n"
    "      AND status = \"ENROLLED\")\n"
    "    IF remainingCount < 2 THEN\n"
    "      Database.Warning.create(\n"
    "        userId = studentId, reason = \"Below minimum course load\")\n"
    "    END IF\n"
    "\n"
    "    RETURN { success: true, action: \"WITHDRAWN\", message: \"Withdrawn with W.\" }\n"
    "\n"
    "  ELSE\n"
    "    RETURN { success: false, message: \"Drop/withdraw not available in current period.\" }\n"
    "  END IF\n"
    "END FUNCTION"
)

doc.add_page_break()

# ════════════════════════════════════════
# UC-19: View & Pay Fine
# ════════════════════════════════════════
heading("UC-19: View and Pay Fine", level=2)

heading("Use-Case Scenario", level=3)
para("Normal Scenario:", bold=True)
bullet("1. Student navigates to the account/financial section of their dashboard.")
bullet("2. System displays current fineOwed balance.")
bullet("3. Student initiates a payment (amount <= fineOwed).")
bullet("4. System validates payment amount and updates user.fineOwed -= payment.")
bullet("5. If fineOwed reaches $0.00, the financial hold is lifted: enrollment and graduation re-enabled.")
bullet("6. System returns confirmation with updated balance.")

para("Exceptional Scenarios:", bold=True)
bullet("Overpayment: Payment amount > fineOwed \u2192 400 \"Payment exceeds balance.\"")
bullet("Attempt to enroll with hold: fineOwed > 0 \u2192 enrollment blocked; redirect to payment page.")
bullet("Attempt to graduate with hold: fineOwed > 0 \u2192 graduation blocked.")
bullet("Registrar manually clears fine: Admin can set fineOwed = 0 directly from registrar panel.")
bullet("No fine owed: Student navigates to payment but fineOwed = 0 \u2192 display \"No outstanding balance.\"")

heading("Sequence Diagram \u2014 UC-19", level=3)
code_block(
    "Student          Portal            API               Database\n"
    "  |                |                 |                  |\n"
    "  |<-- show balance|                 |                  |\n"
    "  |   ($200.00)    |                 |                  |\n"
    "  |                |                 |                  |\n"
    "  |-- pay $200 --->|                 |                  |\n"
    "  |                |-- POST          |                  |\n"
    "  |                |   /fines/pay -->|                  |\n"
    "  |                |                 |-- validate       |\n"
    "  |                |                 |   amount <= fine  |\n"
    "  |                |                 |-- UPDATE user    |\n"
    "  |                |                 |   fineOwed -=    |\n"
    "  |                |                 |   amount ------->|\n"
    "  |                |                 |<-- OK, fine=0 ---|\n"
    "  |                |<-- 200          |                  |\n"
    "  |                |   holdLifted -->|                  |\n"
    "  |<-- \"Balance    |                 |                  |\n"
    "  |   cleared\" ----|                 |                  |\n"
    "  |                                                     |\n"
    "  | [EXCEPTIONAL: overpayment]                          |\n"
    "  |-- pay $500 --->|                 |                  |\n"
    "  |                |-- POST -------->|                  |\n"
    "  |                |                 |-- 500 > 200      |\n"
    "  |                |<-- 400 error ---|                  |\n"
    "  |<-- \"Exceeds\" --|                 |                  |"
)

heading("Pseudo-Code \u2014 payFine(studentId, amount)", level=3)
code_block(
    "FUNCTION payFine(studentId: Int, amount: Float) -> Result\n"
    "  INPUT: studentId -- the student paying\n"
    "         amount -- payment amount\n"
    "  OUTPUT: { success, newBalance, holdLifted?, message }\n"
    "\n"
    "  student = Database.User.findUnique(WHERE id = studentId)\n"
    "\n"
    "  IF student.fineOwed = 0 THEN\n"
    "    RETURN { success: false, message: \"No outstanding balance.\" }\n"
    "  END IF\n"
    "\n"
    "  IF amount <= 0 THEN\n"
    "    RETURN { success: false, message: \"Invalid payment amount.\" }\n"
    "  END IF\n"
    "\n"
    "  IF amount > student.fineOwed THEN\n"
    "    RETURN { success: false, message: \"Payment exceeds balance.\" }\n"
    "  END IF\n"
    "\n"
    "  newBalance = student.fineOwed - amount\n"
    "  Database.User.update(\n"
    "    WHERE id = studentId,\n"
    "    SET fineOwed = newBalance\n"
    "  )\n"
    "\n"
    "  holdLifted = (newBalance = 0)\n"
    "\n"
    "  // If fine fully paid and student was suspended due to fines,\n"
    "  // suspension can now be reviewed by registrar\n"
    "  IF holdLifted AND student.suspended THEN\n"
    "    // Note: registrar must still manually lift suspension\n"
    "    // but financial hold is automatically cleared\n"
    "  END IF\n"
    "\n"
    "  RETURN { success: true, newBalance, holdLifted,\n"
    "           message: holdLifted ? \"Balance cleared. Hold lifted.\"\n"
    "                               : \"Payment applied. Remaining: $\" + newBalance }\n"
    "END FUNCTION"
)

heading("Pseudo-Code \u2014 viewFine(studentId)", level=3)
code_block(
    "FUNCTION viewFine(studentId: Int) -> Result\n"
    "  INPUT: studentId\n"
    "  OUTPUT: { fineOwed, hasHold }\n"
    "\n"
    "  student = Database.User.findUnique(WHERE id = studentId)\n"
    "  RETURN {\n"
    "    fineOwed: student.fineOwed,\n"
    "    hasHold: student.fineOwed > 0\n"
    "  }\n"
    "END FUNCTION"
)

doc.add_page_break()

# ════════════════════════════════════════
# UC-20: Schedule Builder
# ════════════════════════════════════════
heading("UC-20: Schedule Builder", level=2)

heading("Use-Case Scenario", level=3)
para("Normal Scenario:", bold=True)
bullet("1. Student navigates to /schedule_builder during REGISTRATION period.")
bullet("2. System fetches all available (non-cancelled) courses for the current semester with schedule strings.")
bullet("3. Student adds courses to a draft schedule; UI parses schedule strings (e.g., \"MWF 10:00-11:00\") and displays a weekly grid.")
bullet("4. System detects time conflicts in real-time via client-side parsing of schedule strings.")
bullet("5. Conflicting courses are highlighted in red; non-conflicting courses shown in green.")
bullet("6. Student confirms desired courses; system initiates enrollment for each selected course via the registerForCourse function (UC-6).")

para("Exceptional Scenarios:", bold=True)
bullet("Schedule conflict detected: Two selected courses overlap \u2192 UI highlights conflict; prevents enrollment of both.")
bullet("Course becomes full during planning: Enrollment attempt returns error \u2192 UI suggests joining waitlist.")
bullet("Maximum 4 courses: Student tries to add a 5th course \u2192 UI prevents addition with message.")
bullet("Student already enrolled: Course already in student's enrollment \u2192 UI shows it as \"Already Enrolled\" (greyed out).")
bullet("No courses available: Empty course list \u2192 display \"No courses available for this semester.\"")

heading("Sequence Diagram \u2014 UC-20", level=3)
code_block(
    "Student          Browser           Server            Database\n"
    "  |                |                 |                  |\n"
    "  |-- open         |                 |                  |\n"
    "  |   builder ---->|                 |                  |\n"
    "  |                |-- GET           |                  |\n"
    "  |                |   /schedule --->|                  |\n"
    "  |                |                 |-- fetch courses  |\n"
    "  |                |                 |   (current sem,  |\n"
    "  |                |                 |    not cancelled) |\n"
    "  |                |                 |   + student's    |\n"
    "  |                |                 |   enrollments -->|\n"
    "  |                |                 |<-- course list --|\n"
    "  |                |<-- render grid -|                  |\n"
    "  |<-- weekly view |                 |                  |\n"
    "  |                |                 |                  |\n"
    "  |-- add course   |                 |                  |\n"
    "  |   to draft --->|                 |                  |\n"
    "  |                |-- parse schedule strings           |\n"
    "  |                |   (client-side)                    |\n"
    "  |                |-- check conflicts                  |\n"
    "  |<-- show grid   |                 |                  |\n"
    "  |   (green/red)  |                 |                  |\n"
    "  |                |                 |                  |\n"
    "  |-- confirm ---->|                 |                  |\n"
    "  |                |-- POST /register|                  |\n"
    "  |                |   (per course)->|                  |\n"
    "  |                |                 |-- registerFor-   |\n"
    "  |                |                 |   Course() ----->|\n"
    "  |                |                 |<-- enrolled -----|  \n"
    "  |                |<-- results -----|                  |\n"
    "  |<-- summary ----|                 |                  |"
)

heading("Pseudo-Code \u2014 getScheduleData(studentId)", level=3)
code_block(
    "FUNCTION getScheduleData(studentId: Int) -> ScheduleData\n"
    "  INPUT: studentId\n"
    "  OUTPUT: { availableCourses, currentEnrollments, weeklyGrid }\n"
    "\n"
    "  semester = Database.Semester.findFirst(WHERE isCurrent = true)\n"
    "  IF semester IS null OR semester.period != \"REGISTRATION\" THEN\n"
    "    RETURN { error: \"Schedule builder only available during registration.\" }\n"
    "\n"
    "  availableCourses = Database.Course.findMany(\n"
    "    WHERE semesterId = semester.id AND NOT cancelled,\n"
    "    INCLUDE instructor, enrollments (count)\n"
    "  )\n"
    "\n"
    "  currentEnrollments = Database.Enrollment.findMany(\n"
    "    WHERE userId = studentId AND course.semesterId = semester.id,\n"
    "    INCLUDE course\n"
    "  )\n"
    "\n"
    "  RETURN { availableCourses, currentEnrollments }\n"
    "END FUNCTION"
)

heading("Pseudo-Code \u2014 checkScheduleConflict(schedule1, schedule2)", level=3)
code_block(
    "FUNCTION checkScheduleConflict(schedule1: String, schedule2: String) -> Boolean\n"
    "  INPUT: two schedule strings (e.g., \"MWF 10:00-11:00\", \"MW 10:30-11:30\")\n"
    "  OUTPUT: true if conflict exists, false otherwise\n"
    "\n"
    "  days1, startTime1, endTime1 = parseSchedule(schedule1)\n"
    "  days2, startTime2, endTime2 = parseSchedule(schedule2)\n"
    "\n"
    "  // Check if any days overlap\n"
    "  commonDays = INTERSECT(days1, days2)\n"
    "  IF commonDays IS empty THEN\n"
    "    RETURN false  // no overlapping days\n"
    "  END IF\n"
    "\n"
    "  // Check if times overlap on common days\n"
    "  IF startTime1 < endTime2 AND startTime2 < endTime1 THEN\n"
    "    RETURN true   // time overlap exists\n"
    "  END IF\n"
    "\n"
    "  RETURN false\n"
    "END FUNCTION\n"
    "\n"
    "FUNCTION parseSchedule(schedule: String) -> { days[], startTime, endTime }\n"
    "  // Example: \"MWF 10:00-11:00\"\n"
    "  parts = schedule.split(\" \")\n"
    "  dayString = parts[0]  // \"MWF\"\n"
    "  timeRange = parts[1]  // \"10:00-11:00\"\n"
    "\n"
    "  days = []\n"
    "  FOR EACH char IN dayString:\n"
    "    IF char = \"M\" THEN days.push(\"Monday\")\n"
    "    IF char = \"T\" THEN days.push(\"Tuesday\")\n"
    "    IF char = \"W\" THEN days.push(\"Wednesday\")\n"
    "    IF char = \"R\" THEN days.push(\"Thursday\")\n"
    "    IF char = \"F\" THEN days.push(\"Friday\")\n"
    "\n"
    "  [startStr, endStr] = timeRange.split(\"-\")\n"
    "  startTime = parseTime(startStr)\n"
    "  endTime = parseTime(endStr)\n"
    "\n"
    "  RETURN { days, startTime, endTime }\n"
    "END FUNCTION"
)

heading("Pseudo-Code \u2014 enrollFromScheduleBuilder(studentId, courseIds[])", level=3)
code_block(
    "FUNCTION enrollFromScheduleBuilder(studentId: Int, courseIds: Int[]) -> Result[]\n"
    "  INPUT: studentId, array of courseIds to enroll\n"
    "  OUTPUT: array of results per course\n"
    "\n"
    "  IF courseIds.length > 4 THEN\n"
    "    RETURN { success: false, message: \"Maximum 4 courses per semester.\" }\n"
    "\n"
    "  results = []\n"
    "  FOR EACH courseId IN courseIds:\n"
    "    result = CALL registerForCourse(studentId, courseId)  // reuse UC-6 logic\n"
    "    results.push({ courseId, ...result })\n"
    "  END FOR\n"
    "\n"
    "  RETURN results\n"
    "END FUNCTION"
)

doc.add_page_break()

# ════════════════════════════════════════
# Update System Screens section
# ════════════════════════════════════════
heading("5.7 Drop/Withdraw Course Page", level=2)
para("Student view showing enrolled courses with \"Drop\" or \"Withdraw\" buttons depending on "
     "whether the drop window is still open. Each course card displays the course code, name, "
     "and schedule. Drop button is green (within window) or replaced by orange \"Withdraw\" button "
     "(after deadline). A confirmation modal appears before finalizing the action, warning about "
     "transcript implications for withdrawals (W grade) and minimum course load requirements.")

heading("5.8 View and Pay Fine Page", level=2)
para("Student financial section showing current fineOwed balance prominently. If balance > $0, "
     "a payment form appears with an amount field and \"Pay\" button. A progress bar or indicator "
     "shows how much of the fine has been paid. Status messages indicate whether a financial hold "
     "is active (blocking enrollment and graduation). After full payment, a green confirmation "
     "banner displays \"Balance cleared. Hold lifted.\" The registrar view includes an admin "
     "override option to manually clear fines.")

heading("5.9 Schedule Builder Page", level=2)
para("Interactive weekly grid view (Monday through Friday, 8AM to 8PM). Available courses are "
     "listed in a sidebar panel; students drag or click courses to add them to the grid. "
     "Time blocks are color-coded: green for added courses, red highlight for conflicts. "
     "The top bar shows total credits selected and remaining slots (max 4 courses). "
     "A \"Confirm & Enroll\" button at the bottom submits all selected courses for enrollment. "
     "Already-enrolled courses appear as locked grey blocks on the grid.")

# ════════════════════════════════════════
# SAVE
# ════════════════════════════════════════
output_path = "/home/samia/CunyZeroLite/docs/GroupK_CunyZero_PhaseII.docx"
doc.save(output_path)
print(f"Updated report saved to {output_path}")
