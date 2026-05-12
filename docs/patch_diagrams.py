"""Swap the two remaining ASCII diagrams (§1.2 class diagram and §2.1
use-case diagram) for their PNG counterparts in the existing Phase II
docx, without touching pasted screenshots or manual edits."""

import os
import sys
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "GroupK_CunyZero_PhaseII_FINAL.docx")

# Each entry: (distinctive substring that appears ONLY in the ASCII block,
#              png filename to swap in, caption, width in cm)
SWAPS = [
    (
        "<<boundary>>",                              # unique to §1.2 ASCII
        "class_diagram.png",
        "Figure 1.1 — System Collaboration Class Diagram (three-layer architecture)",
        16,
    ),
    (
        "CUNYZEROLITE  --  SYSTEM BOUNDARY",        # unique to §2.1 ASCII
        "usecase_diagram.png",
        "Figure 2.0 — CUNYZeroLite Use-Case Diagram (4 actors × 20 use cases)",
        17,
    ),
]

# Paragraph-text replacements (intro/explanatory paragraphs that mentioned
# the old ASCII diagram or its layout)
TEXT_REPLACEMENTS = [
    (
        "The CUNYZeroLite system is organized around a three-layer architecture:",
        "The CUNYZeroLite system is organized around a three-layer architecture. The top tier is the set of browser-side pages and widgets the user interacts with. The middle tier is the set of server-side actions that validate input, enforce business rules, and orchestrate work. The bottom tier is the persistent store of domain entities accessed through an ORM.",
    ),
    (
        "The Browser sends requests to the Next.js control layer.",
        "A user action in the browser layer issues an HTTP request or Server Action call to the control layer. The control layer performs authentication, input validation, and business-rule enforcement, then reads and writes domain entities through the Prisma ORM. Session identity is carried on an HTTP-only cookie; the AI assistant subsystem first consults a local policy knowledge base and falls back to an external model only when no local answer is available.",
    ),
    (
        "The diagram below is the organizing centerpiece of this report.",
        "The diagram below is the organizing centerpiece of this report. It shows all four primary actors (Visitor, Student, Instructor, Registrar), the twenty use cases inside the CUNYZeroLite system boundary, and the associations linking each actor to the use cases they participate in. For readability the diagram itself carries only the primary actor–use-case associations; secondary actor participation and «include» / «extend» relationships between use cases are listed in full immediately underneath.",
    ),
]


def replace_paragraph_with_image(paragraph, png_path, caption, width_cm):
    """Remove all runs in `paragraph`, then insert a centered image and a
    caption paragraph immediately after it."""
    # Clear all runs
    for r in paragraph.runs:
        r.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add picture to the first run (or a fresh run if none)
    run = paragraph.add_run()
    run.add_picture(png_path, width=Cm(width_cm))

    # Insert caption paragraph right after this one
    new_para = paragraph.insert_paragraph_before("")
    # Actually — insert_paragraph_before inserts BEFORE. We want AFTER.
    # Easier: use the OXML element to insert a sibling after.
    from docx.oxml.ns import qn
    from copy import deepcopy
    from docx.oxml import OxmlElement

    cap_p = OxmlElement("w:p")
    paragraph._element.addnext(cap_p)

    # Wrap it as a Paragraph via docx API
    from docx.text.paragraph import Paragraph
    cap_para = Paragraph(cap_p, paragraph._parent)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.name = "Times New Roman"

    # Remove the unwanted empty paragraph that insert_paragraph_before created
    if new_para._element.getparent() is not None and not new_para.text:
        new_para._element.getparent().remove(new_para._element)


def replace_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    first.text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def main():
    if not os.path.exists(DOCX):
        print(f"ERROR: {DOCX} not found")
        sys.exit(1)

    doc = Document(DOCX)

    # --- Image swaps ---
    ascii_replaced = 0
    for marker, png_name, caption, w_cm in SWAPS:
        png_path = os.path.join(HERE, png_name)
        if not os.path.exists(png_path):
            print(f"SKIP: {png_name} not found")
            continue
        for para in doc.paragraphs:
            if marker in para.text:
                replace_paragraph_with_image(para, png_path, caption, w_cm)
                ascii_replaced += 1
                print(f"  Replaced ASCII block containing '{marker[:40]}...' with {png_name}")
                break
        else:
            print(f"  WARN: no paragraph contained marker '{marker[:40]}...'")

    # --- Text paragraph replacements ---
    text_done = {old: False for old, _ in TEXT_REPLACEMENTS}
    for para in doc.paragraphs:
        for old, new in TEXT_REPLACEMENTS:
            if text_done[old]:
                continue
            if para.text.startswith(old):
                replace_paragraph_text(para, new)
                text_done[old] = True
                break

    doc.save(DOCX)

    text_matched = sum(1 for v in text_done.values() if v)
    print(f"ASCII diagrams replaced: {ascii_replaced}/{len(SWAPS)}")
    print(f"Text paragraphs updated: {text_matched}/{len(TEXT_REPLACEMENTS)}")
    for old, done in text_done.items():
        if not done:
            print(f"  UNMATCHED text prefix: {old[:60]}...")


if __name__ == "__main__":
    main()
