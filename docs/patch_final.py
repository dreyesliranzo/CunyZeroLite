"""Patch the existing Phase II docx in place.

Updates only the specific paragraphs that carry meta-commentary ("PNG",
".py", "rendered by …") with content-focused wording, while preserving
tables, pasted screenshots, images, and any other manual edits.
"""

import os
import sys
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "GroupK_CunyZero_PhaseII_FINAL.docx")

# Each entry: (old_prefix_to_match, new_full_text).
# We match by the first ~60 chars so we don't get confused if a small word changed.
REPLACEMENTS = [
    # §2.3 intro
    (
        "This section presents a dynamic-behavior diagram for each",
        "This section shows, for each of the twenty use cases in §2.2, the dynamic interaction between the actors and the internal components of the system that serve their request. Eighteen use cases are presented as UML sequence diagrams: vertical lifelines represent the participants (actor, UI page, controller, domain service, and database), time flows top-to-bottom, solid arrows represent synchronous calls or actor requests, dashed arrows represent returns, activation bars mark the period a participant is executing, and self-messages capture internal validation or computation steps. Two use cases — UC-11 Process Complaints and UC-15 View Role-Based Dashboard — are shown instead as UML collaboration diagrams, which emphasize the structural relationships among the collaborating objects rather than a strict temporal order; this is the natural choice when the behavior is best understood as a controller coordinating several peers in a fan-out pattern rather than as a single linear conversation. Taken together, the twenty diagrams show how each use-case request flows end-to-end through authentication, validation, domain logic, and persistence, and what each participant contributes to the final outcome.",
    ),
    # §2.4 intro
    (
        "Each of the three Petri-nets below models a representative dynamic flow in the system.",
        "Each of the three Petri-nets below models a representative dynamic flow whose correctness depends on branching, guard conditions, and synchronized state changes — the kind of behavior that a sequence diagram cannot easily express. In these nets, circles denote places (the states or preconditions that can hold), filled black bars denote transitions (events that fire when every input place holds a token), directed arrows denote arcs connecting places to transitions, and a black dot inside a place represents an initial marking. Terminal places are shaded pink for adverse outcomes and green for successful outcomes, making each net's reachable final states immediately visible. Figure 2.21 models UC-6 Register for Courses, capturing capacity checks, automatic waitlist routing when a section is full, and successful enrollment. Figure 2.22 models UC-9 Assign Grades and the downstream academic-standing updates — warnings, suspension, and termination — that fire when a student's grades or GPA cross defined thresholds. Figure 2.23 models UC-8 Review Submission and the taboo-word filter that determines whether a review is published openly or hidden from course listings.",
    ),
    # §5 intro
    (
        "Each subsection below describes a major GUI screen and provides a placeholder slot",
        "This section documents the principal screens of the CUNYZeroLite user interface. Each subsection names a screen and briefly describes the role it plays in the system; the screenshot that follows the description shows the screen as implemented.",
    ),
    # §5 screen descriptions (match by a distinctive opening substring)
    (
        "The login page presents a centered form with CUNYZeroLite branding in navy blue.",
        "The authentication entry point for every user of the system. The user supplies an email and password; the system verifies the credentials against the User store and, on success, either forwards the user to the change-password screen (when the account is flagged as mustChangePassword) or directly to the role-based dashboard.",
    ),
    (
        "Displayed when mustChangePassword = true after first login.",
        "The mandatory password-reset screen shown on first login for any account whose mustChangePassword flag is set. The user supplies and confirms a new password; on success, the flag is cleared and the user is forwarded to their role-based dashboard.",
    ),
    (
        "After login, users see a card grid layout tailored to their role:",
        "The post-login landing page. It surfaces the set of actions a user is permitted to perform, drawn from the use cases assigned to their role — registration, scheduling, grades, reviews, and complaints for students; taught courses, grading, and complaint filing for instructors; semester, course, student, application, complaint, graduation, and taboo-word management for the registrar.",
    ),
    (
        "The student dashboard is the most comprehensive screen.",
        "The student's academic home page. It aggregates the student's standing (GPA, active warnings, outstanding fine balance, honor-roll status), the courses they are currently enrolled in with schedule and grade status, and their completed courses with final grades, and offers direct entry into registration, the schedule builder, review submission, and complaint filing.",
    ),
    (
        "Emerald-themed navigation with \"Instructor Portal\" subtitle.",
        "The instructor's teaching home page. It lists the courses they are assigned to teach for the current semester together with each course's enrolled roster and the average student-review rating that feeds the instructor warning rule, and provides entry into grade assignment and complaint filing.",
    ),
    (
        "Red-themed navigation with shield icon.",
        "The registrar's administrative home page. It summarizes system totals (students, instructors, active courses, active warnings), surfaces the work queues that require registrar action (pending applications, pending complaints, pending graduation requests), and links into the management pages for semesters, courses, students, and taboo words.",
    ),
    (
        "Student view showing enrolled courses with \"Drop\" or \"Withdraw\" buttons",
        "Where a student removes a course from their current enrollment. During the drop window the action is a drop, which deletes the enrollment outright; after the drop deadline it becomes a withdrawal, which leaves a W on the transcript but does not affect GPA. The page enforces the minimum course-load rule before submitting.",
    ),
    (
        "Student financial section showing fineOwed balance prominently.",
        "Where a student views their outstanding fine balance and makes a payment. While a balance is owed, a financial hold blocks registration and grade viewing; once the balance is cleared the hold is lifted automatically. The registrar has an administrative override to clear a fine without payment.",
    ),
    (
        "Interactive weekly grid (Monday-Friday, 8AM-8PM).",
        "An interactive planner for building a semester schedule. The student selects courses from the available-course list, and the builder validates time conflicts, total credits, and the four-course maximum before allowing the final Confirm & Enroll action, which creates the underlying Enrollment records.",
    ),
]


def replace_paragraph_text(paragraph, new_text):
    """Replace paragraph text while preserving the first run's formatting."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    first.text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def main():
    if not os.path.exists(DOCX):
        print(f"ERROR: {DOCX} not found")
        sys.exit(1)

    doc = Document(DOCX)
    remaining = {old: new for old, new in REPLACEMENTS}

    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
        for old in list(remaining.keys()):
            if text.startswith(old):
                replace_paragraph_text(para, remaining[old])
                del remaining[old]
                break

    # Also walk table cells (for the revision-history table etc. — not expected
    # to match, but safe to check)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if not text:
                        continue
                    for old in list(remaining.keys()):
                        if text.startswith(old):
                            replace_paragraph_text(para, remaining[old])
                            del remaining[old]
                            break

    doc.save(DOCX)

    matched = len(REPLACEMENTS) - len(remaining)
    print(f"Patched {matched}/{len(REPLACEMENTS)} paragraphs.")
    if remaining:
        print("UNMATCHED (not found in docx):")
        for k in remaining:
            print(f"  - {k[:70]}...")
    else:
        print("All replacements applied. Screenshots and other edits preserved.")


if __name__ == "__main__":
    main()
