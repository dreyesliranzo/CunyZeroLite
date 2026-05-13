#!/usr/bin/env python3
"""Generate FINAL Phase II Design Report with all 20 UCs properly integrated."""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_header_footer():
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        r1 = hp.add_run("CunyZeroLite"); r1.font.size = Pt(9); r1.font.name = 'Times New Roman'
        hp.add_run("\t\t").font.size = Pt(9)
        r2 = hp.add_run("Version: 2.0"); r2.font.size = Pt(9); r2.font.name = 'Times New Roman'
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Confidential              \u00a9 Group K"); r.font.size = Pt(9); r.font.name = 'Times New Roman'

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0,0,0); run.font.name = 'Times New Roman'
    return h

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = bold; run.italic = italic
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    run = p.add_run(text); run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    return p

def code_block(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.63)
    run = p.add_run(text); run.font.name = 'Courier New'; run.font.size = Pt(9)
    return p

def screenshot_placeholder(label):
    """Adds a bordered 1-cell table sized as a screenshot slot. User pastes image inside."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Paste screenshot here — {label} ]")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = 'Times New Roman'
    for _ in range(8):
        ep = cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    return table

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h); run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(10)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]; cell.text = ""
            run = cell.paragraphs[0].add_run(str(val)); run.font.name = 'Times New Roman'; run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Group K"); run.bold = True; run.font.size = Pt(18); run.font.name = 'Times New Roman'
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CunyZeroLite\nDesign Report\nPhase II"); run.bold = True; run.font.size = Pt(24); run.font.name = 'Times New Roman'
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 2.0"); run.bold = True; run.font.size = Pt(14); run.font.name = 'Times New Roman'
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# REVISION HISTORY
# ══════════════════════════════════════════════════════════════
heading("Revision History", 1)
add_table(["Date","Version","Description","Author"],[
    ["21/04/26","2.0","Phase II Design Report covering system architecture, all twenty use cases with scenarios, collaboration/sequence diagrams, Petri-nets, E-R diagram, detailed pseudo-code, and system screens.","Diego Reyes Liranzo, Daniel Olekszyk, Samia Islam, Maisha Islam, Kyle Gosine"]
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
for item in [
    "1.  Introduction",
    "    1.1  Purpose",
    "    1.2  System Collaboration Class Diagram",
    "2.  Use Cases",
    "    2.1  Use-Case Diagram",
    "    2.2  Use-Case Scenarios (UC-1 through UC-20)",
    "    2.3  Collaboration and Sequence Diagrams",
    "    2.4  Petri-Net Diagrams",
    "3.  E-R Diagram",
    "4.  Detailed Design (Pseudo-Code)",
    "5.  System Screens",
    "6.  Meeting Memos",
    "7.  Repository",
]:
    p = doc.add_paragraph()
    run = p.add_run(item); run.font.name = 'Times New Roman'; run.font.size = Pt(12)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════
heading("1. Introduction", 1)
heading("1.1 Purpose", 2)
para("This Phase II Design Report provides the data structures and logic required to implement the CUNYZeroLite college management system as specified in the Phase I Software Requirements Specification. It contains the Entity-Relationship model, detailed pseudo-code for every method, use-case scenarios with collaboration/sequence diagrams, Petri-nets, and representative system screen descriptions. The implementation will be based entirely on this document.")

heading("1.2 System Collaboration Class Diagram", 2)
para("The CUNYZeroLite system is organized around a three-layer architecture. The top tier is the set of browser-side pages and widgets the user interacts with. The middle tier is the set of server-side actions that validate input, enforce business rules, and orchestrate work. The bottom tier is the persistent store of domain entities accessed through an ORM.")

def _embed_arch(png_name, caption, width_cm=16):
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), png_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = 'Times New Roman'
        doc.add_paragraph()
    else:
        para(f"[Diagram {png_name} not available.]", italic=True)

_embed_arch("class_diagram.png",
            "Figure 1.1 — System Collaboration Class Diagram (three-layer architecture)",
            width_cm=16)

para("A user action in the browser layer issues an HTTP request or Server Action call to the control layer. The control layer performs authentication, input validation, and business-rule enforcement, then reads and writes domain entities through the Prisma ORM. Session identity is carried on an HTTP-only cookie; the AI assistant subsystem first consults a local policy knowledge base and falls back to an external model only when no local answer is available.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. USE CASES — ALL 20
# ══════════════════════════════════════════════════════════════
heading("2. Use Cases", 1)

heading("2.1 Use-Case Diagram", 2)
para("The diagram below is the organizing centerpiece of this report. It shows all four primary actors (Visitor, Student, Instructor, Registrar), the twenty use cases inside the CUNYZeroLite system boundary, and the associations linking each actor to the use cases they participate in. For readability the diagram itself carries only the primary actor–use-case associations; secondary actor participation and «include» / «extend» relationships between use cases are listed in full immediately underneath.")

_embed_arch("usecase_diagram.png",
            "Figure 2.0 — CUNYZeroLite Use-Case Diagram (4 actors × 20 use cases)",
            width_cm=17)
para("Relationships between use cases:", bold=True)
bullet("UC-2  <<include>>  UC-3  — Visitor's application triggers Registrar review.")
bullet("UC-10 <<include>>  UC-11 — Every filed complaint is reviewed by Registrar.")
bullet("UC-12 <<include>>  UC-13 — Every graduation application is reviewed by Registrar.")
bullet("UC-4  <<include>>  UC-17 — Advancing to RUNNING auto-invokes running-period rule enforcement.")
bullet("UC-4  <<include>>  UC-9 evaluation — Advancing to COMPLETED auto-invokes academic-standing evaluation (UC-9 post-conditions).")
bullet("UC-6  <<extend>>   UC-7  — When a course is full, registration extends into the waitlist flow.")
bullet("UC-8  <<include>>  UC-16 — Review filtering consults the Registrar-managed taboo-word list.")
bullet("UC-18 <<extend>>   UC-7  — Dropping an enrollment auto-promotes the first waitlisted student.")
bullet("UC-19 <<extend>>   UC-6, UC-13 — An unpaid fine blocks registration and graduation until cleared.")

para("Actors shared across multiple use cases:", bold=True)
bullet("The System itself is an implicit actor in UC-3 (auto-flagging qualified apps), UC-4 (period transition triggers), UC-6 (conflict/capacity checks), UC-8 (taboo filtering), UC-9 (GPA/honor-roll/termination evaluation), UC-17 (batch evaluation), and UC-20 (conflict detection).")
bullet("OpenAI GPT-4o-mini is an external secondary actor in UC-14.")
doc.add_page_break()

heading("2.2 Use-Case Scenarios", 2)

use_cases = [
    ("UC-1: View Public Homepage", [
        ("Actors:", [
            "Primary: Visitor.",
            "Secondary: Student, Instructor, Registrar (can also view the homepage when logged out).",
        ]),
        ("Normal Scenario:", [
            "1. Visitor navigates to the homepage.",
            "2. System displays general introduction, highest-rated courses, lowest-rated courses, and top-GPA students.",
            "3. Lite AI chat widget is visible in the corner.",
            "4. Visitor reads information or interacts with Lite.",
        ]),
        ("Exceptional Scenarios:", [
            "No data exists: System displays placeholder text ('No courses available yet').",
            "Database unreachable: System shows a generic error page.",
        ]),
    ]),
    ("UC-2: Apply as Student or Instructor", [
        ("Actors:", [
            "Primary: Visitor.",
            "Secondary: Registrar (reviews the resulting application in UC-3); System (creates Application record with status PENDING).",
        ]),
        ("Normal Scenario:", [
            "1. Visitor clicks 'Apply' on the homepage.",
            "2. System presents application form.",
            "3. Visitor selects role (Student or Instructor).",
            "4. For Student: enters prior GPA and written justification.",
            "5. For Instructor: enters name and contact info (no justification required).",
            "6. Visitor submits the form.",
            "7. System creates an Application record with status PENDING and displays confirmation.",
        ]),
        ("Exceptional Scenarios:", [
            "Missing required fields: System highlights empty fields and prevents submission.",
            "Prior GPA out of range (Student): System rejects values outside 0.0-4.0.",
            "Duplicate application: System informs the visitor that a pending application already exists.",
        ]),
    ]),
    ("UC-3: Approve/Reject Applications", [
        ("Actors:", [
            "Primary: Registrar.",
            "Secondary: Visitor (awaiting the decision and receiving credentials on accept); System (auto-flags qualified student applications, generates email/temp password).",
        ]),
        ("Normal Scenario:", [
            "1. Registrar navigates to the Applications page.",
            "2. System displays all PENDING applications.",
            "3. For a Student application with prior GPA > 3.0 and quota not reached, system flags for auto-acceptance.",
            "4. Registrar clicks 'Approve.'",
            "5. System generates university email (first initial + last name + 00@cuny.edu) and temporary password.",
            "6. Application status updates to ACCEPTED; new User record created with mustChangePassword = true.",
        ]),
        ("Exceptional Scenarios:", [
            "Reject a qualified student (GPA > 3.0): System requires registrar to provide written justification.",
            "Program quota reached: System alerts registrar; registrar may still override.",
            "Instructor application rejected: Status set to REJECTED with no justification required.",
        ]),
    ]),
    ("UC-4: Manage Semester Periods", [
        ("Actors:", [
            "Primary: Registrar.",
            "Secondary: Student, Instructor (impacted by each period transition); System (auto-invokes UC-17 rules on RUNNING and academic-standing evaluation on COMPLETED).",
        ]),
        ("Normal Scenario:", [
            "1. Registrar opens Semester Management.",
            "2. System shows current semester and its period.",
            "3. Registrar clicks 'Advance Period.'",
            "4. System transitions: CLASS_SETUP -> REGISTRATION -> RUNNING -> GRADING -> COMPLETED.",
            "5. Period updates; UI reflects the new state.",
        ]),
        ("Exceptional Scenarios:", [
            "Advance to COMPLETED triggers academic evaluation: auto-recalculates GPAs, checks honor roll, issues termination/warnings.",
            "No current semester exists: System prompts registrar to create one first.",
            "Advance out of order: System only allows the next sequential period.",
        ]),
    ]),
    ("UC-5: Set Up Courses", [
        ("Actors:", [
            "Primary: Registrar.",
            "Secondary: Instructor (assigned to course; receives warning if course is later cancelled); System (enforces unique code + CLASS_SETUP period).",
        ]),
        ("Normal Scenario:", [
            "1. Registrar opens Course Management during CLASS_SETUP.",
            "2. Registrar enters course code, name, credits, schedule, max students, and selects an instructor.",
            "3. System creates the Course record linked to the current semester.",
            "4. Course appears in the course list.",
        ]),
        ("Exceptional Scenarios:", [
            "Duplicate course code in same semester: System rejects with error.",
            "Cancel a course: system sets cancelled = true and issues a warning to the instructor.",
            "Semester not in CLASS_SETUP: System disables course creation.",
        ]),
    ]),
    ("UC-6: Register for Courses", [
        ("Actors:", [
            "Primary: Student.",
            "Secondary: Instructor (teaches the target course; sees new roster entry); System (enforces conflict, capacity, retake, suspension, and 2–4-course rules).",
        ]),
        ("Normal Scenario:", [
            "1. Student opens Course Registration during REGISTRATION period.",
            "2. System displays available courses for the current semester.",
            "3. Student selects a course.",
            "4. System checks: no time conflict, course not full, student not suspended/terminated, 2-4 courses.",
            "5. System creates Enrollment record with status ENROLLED.",
        ]),
        ("Exceptional Scenarios:", [
            "Time conflict: System rejects and displays conflicting course.",
            "Course full: System creates Waitlist entry instead and shows waitlist position.",
            "Student suspended or terminated: System denies with message.",
            "Already enrolled in 4 courses: System prevents registration.",
            "Retaking a course not previously failed with F: System blocks re-enrollment.",
        ]),
    ]),
    ("UC-7: Manage Waitlist", [
        ("Actors:", [
            "Primary: Instructor.",
            "Secondary: Student (waitlisted — target of admission); System (re-orders remaining waitlist positions).",
        ]),
        ("Normal Scenario:", [
            "1. Instructor opens their course page.",
            "2. System displays waitlisted students in order of request.",
            "3. Instructor clicks 'Admit' next to a student.",
            "4. System changes Waitlist status to ADMITTED, creates an Enrollment record.",
            "5. Waitlist positions update for remaining students.",
        ]),
        ("Exceptional Scenarios:", [
            "Course at capacity and instructor admits: System allows (instructor override).",
            "Student no longer eligible (suspended since waitlisting): System warns instructor and blocks.",
            "Empty waitlist: System displays 'No students on waitlist.'",
        ]),
    ]),
    ("UC-8: Write and Rate Course Reviews", [
        ("Actors:", [
            "Primary: Student (review author).",
            "Secondary: Instructor (receives automatic warning if course average rating drops below 2.0); Registrar (owns the TabooWord list consulted during filtering, per UC-16); System (performs taboo-word replacement and hide logic).",
        ]),
        ("Normal Scenario:", [
            "1. Student navigates to Course Reviews.",
            "2. System shows enrolled courses (grades not yet posted).",
            "3. Student selects a course, enters 1-5 star rating and optional comment.",
            "4. System checks comment against TabooWord list.",
            "5. Zero taboo words found: review is created and published.",
            "6. If course average rating drops below 2.0, instructor receives automatic warning.",
        ]),
        ("Exceptional Scenarios:", [
            "1-2 taboo words: System replaces words with asterisks, publishes review, issues 1 warning.",
            "3+ taboo words: System hides review entirely and issues 2 warnings.",
            "Grades already posted: System prevents review submission.",
            "Duplicate review: System rejects (unique constraint on authorId + courseId).",
        ]),
    ]),
    ("UC-9: Assign Grades", [
        ("Actors:", [
            "Primary: Instructor.",
            "Secondary: Student (grade recipient; may be terminated, warned, or added to honor roll); Registrar (may question instructor on outlier class GPA; triggers termination/honor-roll at COMPLETED); System (runs academic-standing evaluation on finalization).",
        ]),
        ("Normal Scenario:", [
            "1. Instructor opens Grading page during GRADING period.",
            "2. System lists all enrolled students per course.",
            "3. Instructor assigns a letter grade (A, B, C, D, or F) to each student.",
            "4. System saves grades to Enrollment records.",
            "5. When all students graded, system marks grading complete for that course.",
        ]),
        ("Exceptional Scenarios:", [
            "Not all students graded before period ends: Instructor receives a warning.",
            "Class GPA > 3.5 or < 2.5: Registrar may question instructor and issue warning.",
            "Student fails same course twice: System auto-terminates after finalization.",
            "Student cumulative GPA < 2.0: System auto-terminates.",
            "Student GPA 2.0-2.25: Warning issued; must interview registrar.",
            "Semester GPA > 3.75 or cumulative > 3.5 (after >1 semester): Honor roll; can remove one warning.",
        ]),
    ]),
    ("UC-10: File Complaint", [
        ("Actors:", [
            "Primary: Student OR Instructor (the filer).",
            "Secondary: Student or Instructor (the target named in the complaint); Registrar (resolves the complaint in UC-11 and may warn an instructor filer for unjustified filings).",
        ]),
        ("Normal Scenario:", [
            "1. Student or Instructor navigates to File Complaint.",
            "2. User selects target (a student or instructor), enters description.",
            "3. System creates Complaint record with status PENDING.",
            "4. Confirmation shown.",
        ]),
        ("Exceptional Scenarios:", [
            "Complaint against self: System prevents filing.",
            "Empty description: System requires text before submission.",
        ]),
    ]),
    ("UC-11: Process Complaints", [
        ("Actors:", [
            "Primary: Registrar.",
            "Secondary: Student or Instructor (target — receives warning, de-registration, or no action); Student or Instructor (filer — may be warned back for unjustified filing); System (auto-suspends + applies $100 fine when target hits 3 warnings).",
        ]),
        ("Normal Scenario:", [
            "1. Registrar opens Pending Complaints.",
            "2. System lists all PENDING complaints with filer and target details.",
            "3. Registrar reviews and chooses action: issue warning, de-register student, or dismiss.",
            "4. System updates Complaint status to RESOLVED and records resolution.",
        ]),
        ("Exceptional Scenarios:", [
            "Instructor-filed complaint: Registrar must either punish target or warn instructor for unjustified filing.",
            "Action triggers 3rd warning: System auto-suspends student, sets fineOwed.",
            "Dismiss complaint: Status set to DISMISSED, no action taken.",
        ]),
    ]),
    ("UC-12: Apply for Graduation", [
        ("Actors:", [
            "Primary: Student.",
            "Secondary: Registrar (approver in UC-13); System (counts completed passing courses and flags reckless applications).",
        ]),
        ("Normal Scenario:", [
            "1. Student navigates to Graduation Application.",
            "2. System verifies student has completed 8 courses with passing grades.",
            "3. Student submits application.",
            "4. System creates GraduationRequest with status PENDING.",
        ]),
        ("Exceptional Scenarios:", [
            "Fewer than 8 completed courses: System warns this is premature; student may still submit.",
            "Reckless application: If registrar rejects, student receives a warning.",
        ]),
    ]),
    ("UC-13: Approve/Reject Graduation", [
        ("Actors:", [
            "Primary: Registrar.",
            "Secondary: Student (applicant — graduated on approve, warned on reject if reckless); System (verifies 8 passing courses and no outstanding holds).",
        ]),
        ("Normal Scenario:", [
            "1. Registrar opens Graduation Requests.",
            "2. System displays pending requests with student course history.",
            "3. Registrar verifies 8 courses with passing grades and no outstanding holds.",
            "4. Registrar approves: student's graduated flag set to true; student leaves system.",
        ]),
        ("Exceptional Scenarios:", [
            "Outstanding holds (fines, suspension): Registrar rejects; warning issued.",
            "Fewer than 8 passing courses: Registrar rejects; reckless application warning.",
        ]),
    ]),
    ("UC-14: Ask AI Assistant (Lite)", [
        ("Actors:", [
            "Primary: Any authenticated or anonymous user — Visitor, Student, Instructor, or Registrar.",
            "Secondary: OpenAI GPT-4o-mini (external LLM service); System (performs keyword scoring over policy.json for RAG context).",
        ]),
        ("Normal Scenario:", [
            "1. User clicks the Lite chat widget.",
            "2. User types a question.",
            "3. System searches policy.json for matching keywords (RAG).",
            "4. Match found: relevant policy passed as context to GPT-4o-mini; grounded response returned.",
            "5. No match: GPT-4o-mini responds from general knowledge with hallucination warning.",
        ]),
        ("Exceptional Scenarios:", [
            "OpenAI API key invalid/missing: System returns 'AI assistant is currently unavailable.'",
            "API rate limit exceeded: System shows retry message.",
            "Empty question: System prompts user to enter a question.",
        ]),
    ]),
    ("UC-15: View Role-Based Dashboard", [
        ("Actors:", [
            "Primary: Student, Instructor, or Registrar (any authenticated user).",
            "Secondary: System (reads the session cookie, determines role, and dispatches to the correct dashboard variant).",
        ]),
        ("Normal Scenario:", [
            "1. User logs in successfully.",
            "2. System reads session cookie and determines role.",
            "3. System redirects to appropriate dashboard (Student/Instructor/Registrar).",
            "4. Dashboard data fetched from database and rendered.",
        ]),
        ("Exceptional Scenarios:", [
            "Session expired/invalid: Redirect to login.",
            "First-time login (mustChangePassword = true): Redirect to change-password page.",
            "Account terminated/suspended/fired between sessions: Login check prevents access.",
        ]),
    ]),
    ("UC-16: Manage Taboo Words", [
        ("Actors:", [
            "Primary: Registrar.",
            "Secondary: Student (downstream — reviews authored under UC-8 are filtered or hidden against this list); System (enforces lowercase-uniqueness).",
        ]),
        ("Normal Scenario:", [
            "1. Registrar opens Taboo Words page.",
            "2. System displays current list of taboo words.",
            "3. Registrar adds a new word; system creates TabooWord record.",
            "4. Registrar removes a word; system deletes the record.",
        ]),
        ("Exceptional Scenarios:", [
            "Duplicate word: System rejects (unique constraint).",
            "Empty input: System prevents submission.",
        ]),
    ]),
    ("UC-17: Enforce Class Running Period Rules", [
        ("Actors:", [
            "Primary: Registrar (triggers by advancing semester to RUNNING, per UC-4).",
            "Secondary: Student (warned if enrolled in fewer than 2 courses; may get special registration window if their course is cancelled); Instructor (warned when a course is cancelled; suspended if all of their courses are cancelled); System (performs batch evaluation of all enrollments and courses).",
        ]),
        ("Normal Scenario:", [
            "1. Registrar advances semester to RUNNING.",
            "2. System evaluates: students with < 2 courses get a warning.",
            "3. Courses with < 3 students are cancelled; affected students get special registration window.",
            "4. Instructors of cancelled courses receive a warning.",
        ]),
        ("Exceptional Scenarios:", [
            "Instructor's entire course load cancelled: Instructor suspended, cannot teach next semester.",
            "Student already at 3 warnings after check: Auto-suspension triggered.",
            "No courses meet cancellation criteria: System proceeds without changes.",
        ]),
    ]),
    # ── NEW USE CASES ──
    ("UC-18: Drop / Withdraw from Course", [
        ("Actors:", [
            "Primary: Student.",
            "Secondary: Waitlisted Student (auto-promoted to ENROLLED when a seat opens); Instructor (sees updated roster); System (validates date against drop window, performs waitlist promotion, warns on sub-minimum load).",
        ]),
        ("Normal Scenario \u2014 Drop (within first 2 weeks):", [
            "1. Student clicks 'Drop Course' within the drop window.",
            "2. System verifies current date is within the drop deadline (first 2 weeks of RUNNING).",
            "3. System deletes the Enrollment record entirely (no transcript entry).",
            "4. If students are on the waitlist, the first waitlisted student is auto-promoted to ENROLLED.",
            "5. System returns success; course removed from schedule.",
        ]),
        ("Normal Scenario \u2014 Withdraw (after drop deadline):", [
            "1. Student clicks 'Withdraw from Course' after the drop deadline.",
            "2. System verifies semester is in RUNNING period (not COMPLETED).",
            "3. System updates Enrollment: status = 'WITHDRAWN', grade = 'W'.",
            "4. W appears on transcript but does not affect GPA.",
            "5. System checks if student drops below minimum course load (2 courses) and warns.",
        ]),
        ("Exceptional Scenarios:", [
            "Drop window closed: Date past 2-week deadline \u2192 redirects to Withdraw flow.",
            "Semester COMPLETED: Cannot drop or withdraw \u2192 400 error.",
            "Not enrolled: Enrollment record not found \u2192 404 error.",
            "Below minimum load: Dropping below 2 courses triggers a warning.",
        ]),
    ]),
    ("UC-19: View and Pay Fine", [
        ("Actors:", [
            "Primary: Student.",
            "Secondary: Registrar (may manually clear a fine directly via admin controls); System (validates payment amount, updates balance, lifts registration/graduation hold when balance reaches $0).",
        ]),
        ("Normal Scenario:", [
            "1. Student navigates to the account/financial section.",
            "2. System displays current fineOwed balance.",
            "3. Student initiates a payment (amount <= fineOwed).",
            "4. System validates payment amount and updates user.fineOwed -= payment.",
            "5. If fineOwed reaches $0.00, financial hold is lifted: enrollment and graduation re-enabled.",
            "6. System returns confirmation with updated balance.",
        ]),
        ("Exceptional Scenarios:", [
            "Overpayment: Payment amount > fineOwed \u2192 400 'Payment exceeds balance.'",
            "Attempt to enroll with hold: fineOwed > 0 \u2192 enrollment blocked; redirect to payment.",
            "Registrar manually clears fine: Admin can set fineOwed = 0 directly.",
            "No fine owed: fineOwed = 0 \u2192 display 'No outstanding balance.'",
        ]),
    ]),
    ("UC-20: Schedule Builder", [
        ("Actors:", [
            "Primary: Student.",
            "Secondary: Instructor (course owner — visible on each course card); System (client-side conflict detection; downstream enrollment via UC-6).",
        ]),
        ("Normal Scenario:", [
            "1. Student navigates to /schedule_builder during REGISTRATION period.",
            "2. System fetches all available courses for the current semester with schedule strings.",
            "3. Student adds courses to a draft schedule; UI displays a weekly grid.",
            "4. System detects time conflicts in real-time via client-side parsing.",
            "5. Conflicting courses are highlighted; non-conflicting courses shown as valid.",
            "6. Student confirms desired courses; system enrolls via registerForCourse (UC-6).",
        ]),
        ("Exceptional Scenarios:", [
            "Schedule conflict: Two courses overlap \u2192 UI highlights conflict; prevents dual enrollment.",
            "Course becomes full during planning: Enrollment fails \u2192 UI suggests waitlist.",
            "Maximum 4 courses: Student tries to add a 5th \u2192 UI prevents with message.",
            "Already enrolled: Course shown as 'Already Enrolled' (greyed out).",
        ]),
    ]),
]

for uc_title, scenarios in use_cases:
    heading(uc_title, 3)
    for scenario_title, steps in scenarios:
        para(scenario_title, bold=True)
        for step in steps:
            bullet(step)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2.2 DIAGRAMS
# ══════════════════════════════════════════════════════════════
heading("2.3 Collaboration and Sequence Diagrams", 2)

para("This section shows, for each of the twenty use cases in §2.2, the dynamic interaction between the actors and the internal components of the system that serve their request. Eighteen use cases are presented as UML sequence diagrams: vertical lifelines represent the participants (actor, UI page, controller, domain service, and database), time flows top-to-bottom, solid arrows represent synchronous calls or actor requests, dashed arrows represent returns, activation bars mark the period a participant is executing, and self-messages capture internal validation or computation steps. Two use cases — UC-11 Process Complaints and UC-15 View Role-Based Dashboard — are shown instead as UML collaboration diagrams, which emphasize the structural relationships among the collaborating objects rather than a strict temporal order; this is the natural choice when the behavior is best understood as a controller coordinating several peers in a fan-out pattern rather than as a single linear conversation. Taken together, the twenty diagrams show how each use-case request flows end-to-end through authentication, validation, domain logic, and persistence, and what each participant contributes to the final outcome.")

doc.add_paragraph()

def _embed_seq(png_name, caption, width_cm=15):
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), png_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = 'Times New Roman'
        doc.add_paragraph()
    else:
        para(f"[Diagram {png_name} not available.]", italic=True)

_seq_diagrams = [
    ("seq_uc1.png",  "Figure 2.1 — UC-1 View Public Homepage (Sequence Diagram)"),
    ("seq_uc2.png",  "Figure 2.2 — UC-2 Apply as Student or Instructor (Sequence Diagram)"),
    ("seq_uc3.png",  "Figure 2.3 — UC-3 Approve/Reject Applications (Sequence Diagram)"),
    ("seq_uc4.png",  "Figure 2.4 — UC-4 Manage Semester Periods (Sequence Diagram)"),
    ("seq_uc5.png",  "Figure 2.5 — UC-5 Set Up Courses (Sequence Diagram)"),
    ("seq_uc6.png",  "Figure 2.6 — UC-6 Register for Courses (Sequence Diagram)"),
    ("seq_uc7.png",  "Figure 2.7 — UC-7 Manage Waitlist (Sequence Diagram)"),
    ("seq_uc8.png",  "Figure 2.8 — UC-8 Write and Rate Course Reviews (Sequence Diagram)"),
    ("seq_uc9.png",  "Figure 2.9 — UC-9 Assign Grades & Academic Standing (Sequence Diagram)"),
    ("seq_uc10.png", "Figure 2.10 — UC-10 File Complaint (Sequence Diagram)"),
    ("seq_uc11.png", "Figure 2.11 — UC-11 Process Complaints (Collaboration Diagram)"),
    ("seq_uc12.png", "Figure 2.12 — UC-12 Apply for Graduation (Sequence Diagram)"),
    ("seq_uc13.png", "Figure 2.13 — UC-13 Approve/Reject Graduation (Sequence Diagram)"),
    ("seq_uc14.png", "Figure 2.14 — UC-14 Ask AI Assistant (Lite) (Sequence Diagram)"),
    ("seq_uc15.png", "Figure 2.15 — UC-15 View Role-Based Dashboard (Collaboration Diagram)"),
    ("seq_uc16.png", "Figure 2.16 — UC-16 Manage Taboo Words (Sequence Diagram)"),
    ("seq_uc17.png", "Figure 2.17 — UC-17 Enforce Class Running Period Rules (Sequence Diagram)"),
    ("seq_uc18.png", "Figure 2.18 — UC-18 Drop / Withdraw from Course (Sequence Diagram)"),
    ("seq_uc19.png", "Figure 2.19 — UC-19 View and Pay Fine (Sequence Diagram)"),
    ("seq_uc20.png", "Figure 2.20 — UC-20 Schedule Builder (Sequence Diagram)"),
]

for png, caption in _seq_diagrams:
    _embed_seq(png, caption, width_cm=15)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2.4 PETRI-NETS
# ══════════════════════════════════════════════════════════════
heading("2.4 Petri-Net Diagrams", 2)

para("Each of the three Petri-nets below models a representative dynamic flow whose correctness depends on branching, guard conditions, and synchronized state changes — the kind of behavior that a sequence diagram cannot easily express. In these nets, circles denote places (the states or preconditions that can hold), filled black bars denote transitions (events that fire when every input place holds a token), directed arrows denote arcs connecting places to transitions, and a black dot inside a place represents an initial marking. Terminal places are shaded pink for adverse outcomes and green for successful outcomes, making each net's reachable final states immediately visible. Figure 2.21 models UC-6 Register for Courses, capturing capacity checks, automatic waitlist routing when a section is full, and successful enrollment. Figure 2.22 models UC-9 Assign Grades and the downstream academic-standing updates — warnings, suspension, and termination — that fire when a student's grades or GPA cross defined thresholds. Figure 2.23 models UC-8 Review Submission and the taboo-word filter that determines whether a review is published openly or hidden from course listings.")

doc.add_paragraph()

def _embed_petri(png_name, caption, width_cm=16):
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), png_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = 'Times New Roman'
        doc.add_paragraph()
    else:
        para(f"[Diagram {png_name} not available.]", italic=True)

_embed_petri("petri_uc6.png",
             "Figure 2.21 — Petri-Net for UC-6 Register for Courses",
             width_cm=14)
_embed_petri("petri_uc9.png",
             "Figure 2.22 — Petri-Net for UC-9 Assign Grades & Academic Standing",
             width_cm=15)
_embed_petri("petri_uc8.png",
             "Figure 2.23 — Petri-Net for UC-8 Review Submission & Taboo Filtering",
             width_cm=16)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 3. E-R DIAGRAM
# ══════════════════════════════════════════════════════════════
heading("3. E-R Diagram", 1)
para("The CUNYZeroLite database consists of 12 entities. Figure 3.1 below is the main Chen-notation E-R diagram showing every entity, its primary-key attribute (an underlined oval), 1–3 representative non-primary-key attributes, and all relationships with cardinality labels (1, N, or M). Rectangles denote strong entities; double-bordered rectangles denote weak/associative entities (Enrollment, Waitlist, Review, Complaint, HonorRoll); diamonds denote relationships, with double-outlined diamonds indicating identifying relationships on weak entities. The Review—moderates—TabooWord association is a conceptual M:N relationship enforced at the application layer during review submission (UC-8) rather than by a foreign-key constraint. Figure 3.2, shown immediately after the main diagram, is a separate sub-figure detailing the User specialization: an ISA triangle connects User (the supertype) to its four role subtypes (STUDENT, INSTRUCTOR, REGISTRAR, VISITOR), with the User.role field acting as the discriminator (disjoint and total). Complete attribute lists (including types, defaults, and constraints) are provided in the tables immediately following the two diagrams.")

# Embed the Chen-notation E-R diagram PNGs (main + ISA sub-figure)
def _embed_er(png_name, caption, width_cm=17):
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), png_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = 'Times New Roman'
        doc.add_paragraph()
    else:
        para(f"[Diagram {png_name} not available.]", italic=True)

_embed_er("er_diagram.png",
          "Figure 3.1 — CunyZeroLite Entity-Relationship Diagram (Chen notation)",
          width_cm=17)
_embed_er("er_diagram_isa.png",
          "Figure 3.2 — User Specialization (ISA Hierarchy)",
          width_cm=13)

doc.add_paragraph()
para("The tables below list every entity with its attributes, types, primary keys, foreign keys, and constraints in full detail.", bold=True)

entities = [
    ("User", [["id","Integer","PK, auto-increment"],["email","String","UNIQUE"],["username","String","UNIQUE"],["password","String",""],["firstName","String",""],["lastName","String",""],["role","String","Default: STUDENT"],["gpa","Float","Default: 0.0"],["warnings","Int","Default: 0"],["suspended","Boolean","Default: false"],["terminated","Boolean","Default: false"],["fired","Boolean","Default: false"],["graduated","Boolean","Default: false"],["fineOwed","Float","Default: 0.0"],["mustChangePassword","Boolean","Default: true"],["createdAt","DateTime","Auto"],["updatedAt","DateTime","Auto"]]),
    ("Semester", [["id","Integer","PK, auto-increment"],["name","String","UNIQUE"],["year","Int",""],["term","String",""],["period","String","Default: CLASS_SETUP"],["startDate","DateTime",""],["endDate","DateTime",""],["isCurrent","Boolean","Default: false"],["programQuota","Int","Default: 50"],["createdAt","DateTime","Auto"]]),
    ("Course", [["id","Integer","PK, auto-increment"],["code","String","UNIQUE(code, semesterId)"],["name","String",""],["credits","Int","Default: 3"],["maxStudents","Int","Default: 30"],["schedule","String",""],["cancelled","Boolean","Default: false"],["createdAt","DateTime","Auto"],["semesterId","Int","FK -> Semester.id"],["instructorId","Int?","FK -> User.id (nullable)"]]),
    ("Enrollment", [["id","Integer","PK, auto-increment"],["status","String","Default: ENROLLED"],["grade","String?","Nullable"],["createdAt","DateTime","Auto"],["userId","Int","FK -> User.id, UNIQUE(userId,courseId)"],["courseId","Int","FK -> Course.id"]]),
    ("Waitlist", [["id","Integer","PK, auto-increment"],["status","String","Default: WAITING"],["position","Int",""],["createdAt","DateTime","Auto"],["userId","Int","FK -> User.id, UNIQUE(userId,courseId)"],["courseId","Int","FK -> Course.id"]]),
    ("Review", [["id","Integer","PK, auto-increment"],["rating","Int","1-5"],["comment","String?","Nullable"],["hidden","Boolean","Default: false"],["createdAt","DateTime","Auto"],["authorId","Int","FK -> User.id, UNIQUE(authorId,courseId)"],["courseId","Int","FK -> Course.id"]]),
    ("Complaint", [["id","Integer","PK, auto-increment"],["description","String",""],["status","String","Default: PENDING"],["resolution","String?","Nullable"],["createdAt","DateTime","Auto"],["filerId","Int","FK -> User.id"],["targetId","Int","FK -> User.id"]]),
    ("Warning", [["id","Integer","PK, auto-increment"],["reason","String",""],["removed","Boolean","Default: false"],["createdAt","DateTime","Auto"],["userId","Int","FK -> User.id"]]),
    ("Application", [["id","Integer","PK, auto-increment"],["type","String","STUDENT | INSTRUCTOR"],["status","String","Default: PENDING"],["priorGpa","Float?","Nullable"],["justification","String?","Nullable"],["createdAt","DateTime","Auto"],["userId","Int","FK -> User.id"]]),
    ("GraduationRequest", [["id","Integer","PK, auto-increment"],["status","String","Default: PENDING"],["createdAt","DateTime","Auto"],["userId","Int","FK -> User.id"]]),
    ("TabooWord", [["id","Integer","PK, auto-increment"],["word","String","UNIQUE"]]),
    ("HonorRoll", [["id","Integer","PK, auto-increment"],["type","String","SEMESTER | OVERALL"],["usedToRemoveWarning","Boolean","Default: false"],["createdAt","DateTime","Auto"],["userId","Int","FK -> User.id, UNIQUE(userId,semesterId,type)"],["semesterId","Int","FK -> Semester.id"]]),
]

for name, attrs in entities:
    heading(name, 3)
    add_table(["Attribute","Type","Constraint"], attrs)
    doc.add_paragraph()

heading("Relationship Summary", 3)
for r in [
    "User (1) --- (M) Course: instructor teaches courses",
    "User (M) --- (M) Course via Enrollment: students enroll in courses",
    "User (M) --- (M) Course via Waitlist: students waitlisted for courses",
    "User (M) --- (M) Course via Review: students review courses",
    "User (1) --- (M) Warning: user receives warnings",
    "User (1) --- (M) Application: user submits applications",
    "User (1) --- (M) GraduationRequest: student applies for graduation",
    "User (M) --- (M) User via Complaint: filer complains about target",
    "User (M) --- (M) Semester via HonorRoll: student honored per semester",
    "Semester (1) --- (M) Course: semester contains courses",
    "Review (M) --- (N) TabooWord via moderates: taboo words filter / hide reviews at application layer (no FK; enforced during UC-8 submission)",
]: bullet(r)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. DETAILED DESIGN
# ══════════════════════════════════════════════════════════════
heading("4. Detailed Design", 1)
para("This section provides pseudo-code for every method in the system, organized by subsystem.")

pseudocode = [
    ("4.1 Session Management", [
        ("createSession(userId)", "FUNCTION createSession(userId: Integer) -> SessionData | null\n  INPUT: userId -- the database ID of the authenticated user\n  OUTPUT: SessionData object or null if user not found\n\n  user = Database.User.findUnique(WHERE id = userId)\n  IF user IS null THEN RETURN null\n\n  session = { userId: user.id, role: user.role,\n    firstName: user.firstName, lastName: user.lastName, email: user.email }\n\n  SET httpOnly cookie \"session\" = JSON.stringify(session)\n    WITH maxAge = 86400 seconds, sameSite = \"lax\", path = \"/\"\n\n  RETURN session\nEND FUNCTION"),
        ("getSession()", "FUNCTION getSession() -> SessionData | null\n  INPUT: none (reads from HTTP cookies)\n  OUTPUT: SessionData object or null\n\n  raw = READ cookie \"session\"\n  IF raw IS empty THEN RETURN null\n  TRY: RETURN JSON.parse(raw) AS SessionData\n  CATCH: RETURN null\nEND FUNCTION"),
        ("destroySession()", "FUNCTION destroySession() -> void\n  DELETE cookie \"session\"\nEND FUNCTION"),
    ]),
    ("4.2 Authentication", [
        ("loginUser(email, password)", "FUNCTION loginUser(email, password) -> LoginResult\n  user = Database.User.findUnique(WHERE email = email.trim().toLowerCase())\n  IF user IS null OR user.password != password THEN\n    RETURN { success: false, error: \"Invalid email or password.\" }\n  IF user.terminated THEN RETURN { success: false, error: \"Account terminated.\" }\n  IF user.suspended THEN RETURN { success: false, error: \"Account suspended.\" }\n  IF user.fired THEN RETURN { success: false, error: \"Account deactivated.\" }\n\n  CALL createSession(user.id)\n  redirect = user.mustChangePassword ? \"/change-password\" : \"/dashboard\"\n  RETURN { success: true, role: user.role, firstName: user.firstName, redirect }\nEND FUNCTION"),
        ("logoutUser()", "FUNCTION logoutUser() -> void\n  CALL destroySession()\nEND FUNCTION"),
        ("changePassword(request)", "FUNCTION changePassword(request: HTTP POST) -> JSON\n  session = CALL getSession()\n  IF session IS null THEN RETURN HTTP 401\n  IF newPassword IS empty OR length < 6 THEN RETURN error\n\n  Database.User.update(WHERE id = session.userId,\n    SET password = newPassword, mustChangePassword = false)\n  CALL createSession(session.userId)\n  RETURN { success: true, redirect: \"/dashboard\" }\nEND FUNCTION"),
    ]),
    ("4.3 Dashboard Routing", [
        ("Dashboard() -- main router", "FUNCTION Dashboard() -> HTML Page\n  session = CALL getSession()\n  IF session IS null THEN REDIRECT to \"/login\"\n\n  SWITCH session.role\n    CASE \"REGISTRAR\": cards = registrarCards (8 cards)\n    CASE \"INSTRUCTOR\": cards = instructorCards (5 cards)\n    DEFAULT: cards = studentCards (6 cards)\n  RENDER page with nav bar, welcome, card grid\nEND FUNCTION"),
        ("StudentDashboard()", "FUNCTION StudentDashboard() -> HTML Page\n  session = CALL getSession()\n  IF session IS null OR role != \"STUDENT\" THEN REDIRECT \"/login\"\n\n  user = Database.User.findUnique(WHERE id = session.userId,\n    INCLUDE enrollments->course->instructor,semester,\n    INCLUDE warningsReceived(removed=false), honorRollEntries)\n  currentSemester = Database.Semester.findFirst(WHERE isCurrent=true)\n  currentEnrollments = FILTER WHERE semester.isCurrent\n  pastEnrollments = FILTER WHERE NOT semester.isCurrent\n\n  RENDER: nav, welcome banner, stats (GPA, courses, warnings, fines),\n    honor roll badge, current courses, past courses, sidebar\nEND FUNCTION"),
        ("InstructorDashboard()", "FUNCTION InstructorDashboard() -> HTML Page\n  session = CALL getSession()\n  IF session IS null OR role != \"INSTRUCTOR\" THEN REDIRECT \"/login\"\n\n  user with warningsReceived; courses with enrollments->user, reviews\n  totalStudents = SUM(enrollments.length)\n  FOR EACH course: avgRating = mean(reviews.rating) or null\n\n  RENDER: nav, welcome, stats, course cards with student lists, warnings\nEND FUNCTION"),
        ("RegistrarDashboard()", "FUNCTION RegistrarDashboard() -> HTML Page\n  session = CALL getSession()\n  IF session IS null OR role != \"REGISTRAR\" THEN REDIRECT \"/login\"\n\n  PARALLEL FETCH: totalStudents, totalInstructors, totalCourses,\n    pendingApplications, pendingComplaints, pendingGraduations,\n    currentSemester, activeWarnings, suspendedStudents\n  topStudents = top 5 by GPA\n\n  RENDER: nav, welcome, stats, pending items, management grid,\n    top students sidebar, suspended alert\nEND FUNCTION"),
    ]),
    ("4.4 Course Registration", [
        ("registerForCourse(studentId, courseId)", "FUNCTION registerForCourse(studentId, courseId) -> Result\n  student = Database.User.findUnique(WHERE id = studentId)\n  IF student.suspended OR student.terminated THEN RETURN error\n\n  course = Database.Course.findUnique(WHERE id=courseId, INCLUDE enrollments,semester)\n  IF course.semester.period != \"REGISTRATION\" THEN RETURN error\n\n  currentCount = COUNT enrollments WHERE userId=studentId AND semester.isCurrent\n  IF currentCount >= 4 THEN RETURN error \"Maximum 4 courses\"\n\n  FOR EACH existing enrollment:\n    IF schedule OVERLAPS THEN RETURN error \"Time conflict\"\n\n  IF previous enrollment exists with non-F grade THEN RETURN error\n\n  IF enrollments.length >= maxStudents THEN\n    CREATE Waitlist entry with next position\n    RETURN { success: true, waitlisted: true }\n\n  CREATE Enrollment(userId, courseId, status=ENROLLED)\n  RETURN { success: true, message: \"Enrolled.\" }\nEND FUNCTION"),
    ]),
    ("4.5 Grade Assignment", [
        ("assignGrades(courseId, grades)", "FUNCTION assignGrades(courseId, grades[]) -> Result\n  session = CALL getSession()\n  IF role != INSTRUCTOR THEN RETURN unauthorized\n  course = findUnique(WHERE id=courseId AND instructorId=session.userId)\n  IF course.semester.period != \"GRADING\" THEN RETURN error\n\n  FOR EACH { studentId, grade } IN grades:\n    IF grade NOT IN [A,B,C,D,F] THEN RETURN error\n    UPDATE Enrollment SET grade=grade, status=COMPLETED\n  RETURN { success: true }\nEND FUNCTION"),
    ]),
    ("4.6 Academic Standing Evaluation", [
        ("evaluateAcademicStanding(semesterId)", "FUNCTION evaluateAcademicStanding(semesterId) -> void\n  gradePoints = { A:4.0, B:3.0, C:2.0, D:1.0, F:0.0 }\n\n  FOR EACH active student:\n    Calculate semesterGPA and cumulativeGPA\n    UPDATE user.gpa = cumulativeGPA\n\n    IF any course failed twice (same code) THEN SET terminated=true; CONTINUE\n    IF cumulativeGPA < 2.0 THEN SET terminated=true; CONTINUE\n    IF cumulativeGPA BETWEEN 2.0 AND 2.25 THEN CREATE Warning\n\n    IF semesterGPA > 3.75 THEN CREATE HonorRoll(type=SEMESTER)\n    IF cumulativeGPA > 3.5 AND completedSemesters > 1 THEN CREATE HonorRoll(type=OVERALL)\n\n    FOR EACH unused honor:\n      IF active warning exists THEN remove warning, mark honor used\n\n    IF warnings >= 3 THEN SET suspended=true, fineOwed += 100\nEND FUNCTION"),
    ]),
    ("4.7 Course Review Submission", [
        ("submitReview(authorId, courseId, rating, comment)", "FUNCTION submitReview(authorId, courseId, rating, comment) -> Result\n  Verify enrollment exists AND grade IS null\n  Check no duplicate review\n\n  tabooWords = Database.TabooWord.findMany()\n  tabooCount = 0\n  FOR EACH taboo: IF comment CONTAINS taboo.word THEN tabooCount++; REPLACE WITH \"***\"\n\n  IF tabooCount >= 3: hidden=true, warnings=2\n  ELSE IF tabooCount >= 1: warnings=1\n  ELSE: warnings=0\n\n  CREATE Review(authorId, courseId, rating, comment, hidden)\n  IF warnings > 0: CREATE Warning records, INCREMENT user.warnings\n\n  avgRating = mean of visible reviews\n  IF avgRating < 2.0 THEN CREATE Warning for instructor\n  RETURN { success: true, warningsIssued }\nEND FUNCTION"),
    ]),
    ("4.8 Complaint Management", [
        ("fileComplaint(filerId, targetId, description)", "FUNCTION fileComplaint(filerId, targetId, description) -> Result\n  IF filerId = targetId THEN RETURN error\n  IF description IS empty THEN RETURN error\n  CREATE Complaint(filerId, targetId, description, status=PENDING)\n  RETURN { success: true }\nEND FUNCTION"),
        ("processComplaint(complaintId, action, resolution)", "FUNCTION processComplaint(complaintId, action, resolution) -> Result\n  complaint = findUnique(complaintId, INCLUDE filer, target)\n\n  IF action = DISMISS: UPDATE status=DISMISSED; RETURN\n  IF action = WARN: CREATE Warning for target; INCREMENT warnings\n  IF action = DEREGISTER: DELETE target's current enrollments; CREATE Warning\n  IF action = WARN_FILER: CREATE Warning for filer (unjustified)\n\n  UPDATE complaint status=RESOLVED\n  IF target.warnings >= 3 AND NOT suspended THEN SET suspended=true, fineOwed += 100\nEND FUNCTION"),
    ]),
    ("4.9 Application Management", [
        ("submitApplication(userId, type, priorGpa, justification)", "FUNCTION submitApplication(userId, type, priorGpa, justification) -> Result\n  IF type=STUDENT AND (priorGpa null OR justification empty) THEN RETURN error\n  IF pending application exists THEN RETURN error\n  CREATE Application(userId, type, status=PENDING, priorGpa, justification)\n  RETURN { success: true }\nEND FUNCTION"),
        ("reviewApplication(appId, decision, justification)", "FUNCTION reviewApplication(appId, decision, justification) -> Result\n  IF decision = REJECT:\n    IF app.type=STUDENT AND priorGpa>3.0 AND no justification THEN RETURN error\n    UPDATE status=REJECTED; RETURN\n\n  // ACCEPT\n  IF type = STUDENT:\n    email = firstName[0] + lastName + \"00@cuny.edu\"\n    UPDATE User SET role=STUDENT, email, password=temp, mustChangePassword=true\n  ELSE: UPDATE User SET role=INSTRUCTOR\n  UPDATE Application status=ACCEPTED\nEND FUNCTION"),
    ]),
    ("4.10 Graduation", [
        ("applyForGraduation(studentId)", "FUNCTION applyForGraduation(studentId) -> Result\n  completedCourses = COUNT Enrollment WHERE grade NOT null AND grade != F\n  isReckless = completedCourses < 8\n  CREATE GraduationRequest(userId=studentId, status=PENDING)\n  IF isReckless THEN RETURN { success: true, warning: \"Fewer than 8 courses\" }\n  RETURN { success: true }\nEND FUNCTION"),
        ("reviewGraduation(requestId, decision)", "FUNCTION reviewGraduation(requestId, decision) -> Result\n  IF decision = APPROVE:\n    IF completedPassing < 8 OR suspended OR fineOwed > 0 THEN RETURN error\n    SET graduated=true; UPDATE status=APPROVED\n  IF decision = REJECT:\n    UPDATE status=REJECTED\n    IF completedPassing < 8 THEN CREATE Warning(\"Reckless application\")\nEND FUNCTION"),
    ]),
    ("4.11 Semester and Running Rules", [
        ("advanceSemesterPeriod(semesterId)", "FUNCTION advanceSemesterPeriod(semesterId) -> Result\n  order = [CLASS_SETUP, REGISTRATION, RUNNING, GRADING, COMPLETED]\n  currentIndex = order.indexOf(semester.period)\n  IF currentIndex = last THEN RETURN error\n  newPeriod = order[currentIndex + 1]\n  UPDATE semester SET period = newPeriod\n  IF newPeriod = RUNNING THEN CALL enforceRunningPeriodRules()\n  IF newPeriod = COMPLETED THEN CALL evaluateAcademicStanding()\nEND FUNCTION"),
        ("enforceRunningPeriodRules(semesterId)", "FUNCTION enforceRunningPeriodRules(semesterId) -> void\n  // 1. Warn students with < 2 courses\n  FOR EACH active student:\n    IF enrollmentCount > 0 AND enrollmentCount < 2 THEN CREATE Warning\n\n  // 2. Cancel courses with < 3 students\n  FOR EACH non-cancelled course:\n    IF enrollments.length < 3 THEN\n      SET cancelled=true; CREATE Warning for instructor\n      DELETE enrollments; flag students for special registration\n\n  // 3. Suspend instructors who lost ALL courses\n  FOR EACH instructor:\n    IF had courses but all cancelled THEN SET suspended=true\nEND FUNCTION"),
    ]),
    ("4.12 Taboo Word Management", [
        ("addTabooWord(word)", "FUNCTION addTabooWord(word) -> Result\n  IF word IS empty THEN RETURN error\n  IF word already exists THEN RETURN error\n  CREATE TabooWord(word = word.toLowerCase())\n  RETURN { success: true }\nEND FUNCTION"),
        ("removeTabooWord(wordId)", "FUNCTION removeTabooWord(wordId) -> Result\n  DELETE TabooWord WHERE id = wordId\n  RETURN { success: true }\nEND FUNCTION"),
    ]),
    ("4.13 Waitlist Management", [
        ("admitFromWaitlist(courseId, studentId)", "FUNCTION admitFromWaitlist(courseId, studentId) -> Result\n  entry = Waitlist.findUnique(userId+courseId)\n  IF not found OR status != WAITING THEN RETURN error\n  student = User.findUnique(studentId)\n  IF suspended OR terminated THEN RETURN error\n\n  CREATE Enrollment(userId, courseId, status=ENROLLED)\n  UPDATE Waitlist SET status=ADMITTED\n\n  remaining = findMany(courseId, status=WAITING, ORDER BY position)\n  FOR i = 0 TO remaining.length-1: UPDATE position = i + 1\nEND FUNCTION"),
    ]),
    ("4.14 Course Setup", [
        ("createCourse(...)", "FUNCTION createCourse(code, name, credits, schedule, maxStudents,\n                      instructorId, semesterId) -> Result\n  IF semester.period != CLASS_SETUP THEN RETURN error\n  IF course code exists in same semester THEN RETURN error\n  CREATE Course(code, name, credits, schedule, maxStudents, instructorId, semesterId)\n  RETURN { success: true, course }\nEND FUNCTION"),
        ("cancelCourse(courseId)", "FUNCTION cancelCourse(courseId) -> Result\n  UPDATE Course SET cancelled=true\n  IF instructorId exists THEN CREATE Warning; INCREMENT warnings\n  RETURN { success: true }\nEND FUNCTION"),
    ]),
    ("4.15 AI Assistant (Lite)", [
        ("askAI(question)", "FUNCTION askAI(question) -> Result\n  IF question IS empty THEN RETURN \"Please enter a question.\"\n\n  // RAG: search local knowledge base\n  policies = LOAD \"data/policy.json\"\n  keywords = TOKENIZE(question.toLowerCase())\n  matchedPolicies = []\n  FOR EACH policy: score = COUNT matching keywords\n    IF score > 0 THEN matchedPolicies.push({ policy, score })\n  SORT by score DESC\n\n  IF matches found THEN\n    context = TOP 3 policies; source = \"knowledge_base\"\n  ELSE prompt = question; source = \"general_llm\"\n\n  TRY: response = OpenAI.chat(model=\"gpt-4o-mini\", prompt)\n  CATCH: answer = \"AI assistant unavailable.\"; source = \"error\"\n  RETURN { answer, source }\nEND FUNCTION"),
    ]),
    # ── NEW PSEUDO-CODE ──
    ("4.16 Drop / Withdraw from Course", [
        ("dropOrWithdrawCourse(studentId, courseId)", "FUNCTION dropOrWithdrawCourse(studentId, courseId) -> Result\n  enrollment = Database.Enrollment.findUnique(WHERE userId=studentId AND courseId=courseId)\n  IF enrollment IS null THEN RETURN error \"Not enrolled.\"\n\n  semester = Database.Semester.findFirst(WHERE isCurrent = true)\n  IF semester.period = \"COMPLETED\" THEN RETURN error \"Cannot modify after semester end.\"\n\n  dropDeadline = semester.startDate + 14 days\n\n  IF currentDate <= dropDeadline AND semester.period = \"RUNNING\" THEN\n    // DROP: delete enrollment entirely\n    Database.Enrollment.delete(WHERE id = enrollment.id)\n    // Promote first waitlisted student\n    nextWaitlist = Waitlist.findFirst(WHERE courseId, status=WAITING, ORDER BY position)\n    IF nextWaitlist EXISTS THEN\n      CREATE Enrollment for waitlisted student; UPDATE Waitlist status=ADMITTED\n    RETURN { success: true, action: \"DROPPED\" }\n\n  ELSE IF semester.period = \"RUNNING\" THEN\n    // WITHDRAW: keep record with W grade\n    UPDATE Enrollment SET status=\"WITHDRAWN\", grade=\"W\"\n    IF remaining enrolled courses < 2 THEN CREATE Warning\n    RETURN { success: true, action: \"WITHDRAWN\" }\n\n  ELSE RETURN error \"Not available in current period.\"\nEND FUNCTION"),
    ]),
    ("4.17 Fine Payment", [
        ("payFine(studentId, amount)", "FUNCTION payFine(studentId, amount) -> Result\n  student = Database.User.findUnique(WHERE id = studentId)\n  IF student.fineOwed = 0 THEN RETURN error \"No outstanding balance.\"\n  IF amount <= 0 THEN RETURN error \"Invalid amount.\"\n  IF amount > student.fineOwed THEN RETURN error \"Payment exceeds balance.\"\n\n  newBalance = student.fineOwed - amount\n  Database.User.update(WHERE id = studentId, SET fineOwed = newBalance)\n  holdLifted = (newBalance = 0)\n\n  RETURN { success: true, newBalance, holdLifted }\nEND FUNCTION"),
        ("viewFine(studentId)", "FUNCTION viewFine(studentId) -> Result\n  student = Database.User.findUnique(WHERE id = studentId)\n  RETURN { fineOwed: student.fineOwed, hasHold: student.fineOwed > 0 }\nEND FUNCTION"),
    ]),
    ("4.18 Schedule Builder", [
        ("getScheduleData(studentId)", "FUNCTION getScheduleData(studentId) -> ScheduleData\n  semester = Database.Semester.findFirst(WHERE isCurrent = true)\n  IF semester IS null OR period != \"REGISTRATION\" THEN RETURN error\n\n  availableCourses = Database.Course.findMany(\n    WHERE semesterId = semester.id AND NOT cancelled,\n    INCLUDE instructor, enrollments (count))\n  currentEnrollments = Database.Enrollment.findMany(\n    WHERE userId = studentId AND course.semesterId = semester.id, INCLUDE course)\n\n  RETURN { availableCourses, currentEnrollments }\nEND FUNCTION"),
        ("checkScheduleConflict(schedule1, schedule2)", "FUNCTION checkScheduleConflict(schedule1, schedule2) -> Boolean\n  days1, start1, end1 = parseSchedule(schedule1)\n  days2, start2, end2 = parseSchedule(schedule2)\n\n  commonDays = INTERSECT(days1, days2)\n  IF commonDays IS empty THEN RETURN false\n  IF start1 < end2 AND start2 < end1 THEN RETURN true\n  RETURN false\nEND FUNCTION\n\nFUNCTION parseSchedule(schedule) -> { days[], startTime, endTime }\n  // Example: \"MWF 10:00-11:00\"\n  parts = schedule.split(\" \")\n  dayString = parts[0]; timeRange = parts[1]\n  days = [M->Monday, T->Tuesday, W->Wednesday, R->Thursday, F->Friday]\n  [startStr, endStr] = timeRange.split(\"-\")\n  RETURN { days, parseTime(startStr), parseTime(endStr) }\nEND FUNCTION"),
        ("enrollFromScheduleBuilder(studentId, courseIds[])", "FUNCTION enrollFromScheduleBuilder(studentId, courseIds[]) -> Result[]\n  IF courseIds.length > 4 THEN RETURN error \"Maximum 4 courses.\"\n  results = []\n  FOR EACH courseId IN courseIds:\n    result = CALL registerForCourse(studentId, courseId)  // reuse UC-6\n    results.push({ courseId, ...result })\n  RETURN results\nEND FUNCTION"),
    ]),
]

for section_title, methods in pseudocode:
    heading(section_title, 2)
    for method_name, pseudo in methods:
        heading(method_name, 3)
        code_block(pseudo)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. SYSTEM SCREENS
# ══════════════════════════════════════════════════════════════
heading("5. System Screens", 1)

screens = [
    ("5.1 Login Page", "The authentication entry point for every user of the system. The user supplies an email and password; the system verifies the credentials against the User store and, on success, either forwards the user to the change-password screen (when the account is flagged as mustChangePassword) or directly to the role-based dashboard."),
    ("5.2 Change Password Page", "The mandatory password-reset screen shown on first login for any account whose mustChangePassword flag is set. The user supplies and confirms a new password; on success, the flag is cleared and the user is forwarded to their role-based dashboard."),
    ("5.3 Main Dashboard (Card Grid)", "The post-login landing page. It surfaces the set of actions a user is permitted to perform, drawn from the use cases assigned to their role — registration, scheduling, grades, reviews, and complaints for students; taught courses, grading, and complaint filing for instructors; semester, course, student, application, complaint, graduation, and taboo-word management for the registrar."),
    ("5.4 Student Dashboard (Sample Prototype)", "The student's academic home page. It aggregates the student's standing (GPA, active warnings, outstanding fine balance, honor-roll status), the courses they are currently enrolled in with schedule and grade status, and their completed courses with final grades, and offers direct entry into registration, the schedule builder, review submission, and complaint filing."),
    ("5.5 Instructor Dashboard", "The instructor's teaching home page. It lists the courses they are assigned to teach for the current semester together with each course's enrolled roster and the average student-review rating that feeds the instructor warning rule, and provides entry into grade assignment and complaint filing."),
    ("5.6 Registrar Dashboard", "The registrar's administrative home page. It summarizes system totals (students, instructors, active courses, active warnings), surfaces the work queues that require registrar action (pending applications, pending complaints, pending graduation requests), and links into the management pages for semesters, courses, students, and taboo words."),
    ("5.7 Drop/Withdraw Course Page", "Where a student removes a course from their current enrollment. During the drop window the action is a drop, which deletes the enrollment outright; after the drop deadline it becomes a withdrawal, which leaves a W on the transcript but does not affect GPA. The page enforces the minimum course-load rule before submitting."),
    ("5.8 View and Pay Fine Page", "Where a student views their outstanding fine balance and makes a payment. While a balance is owed, a financial hold blocks registration and grade viewing; once the balance is cleared the hold is lifted automatically. The registrar has an administrative override to clear a fine without payment."),
    ("5.9 Schedule Builder Page", "An interactive planner for building a semester schedule. The student selects courses from the available-course list, and the builder validates time conflicts, total credits, and the four-course maximum before allowing the final Confirm & Enroll action, which creates the underlying Enrollment records."),
]

para("This section documents the principal screens of the CUNYZeroLite user interface. Each subsection names a screen and briefly describes the role it plays in the system; the screenshot that follows the description shows the screen as implemented.")

for title, desc in screens:
    heading(title, 2)
    para(desc)
    para("Screenshot:", bold=True)
    # Strip leading "5.x " from title to get a clean caption for the placeholder
    caption = title.split(" ", 1)[1] if " " in title else title
    screenshot_placeholder(caption)
    doc.add_paragraph()  # breathing room

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. MEETING MEMOS
# ══════════════════════════════════════════════════════════════
heading("6. Meeting Memos", 1)

meetings = [
    ("Meeting 1 \u2014 Early March 2026",
     "The team held its first meeting to review the project specification and establish the foundation for development. All use cases were identified and discussed. The technology stack was selected (Next.js 16, TypeScript, Prisma 7.5, SQLite, Tailwind CSS 4) based on the project requirements. Team members were assigned to different areas of the system including frontend pages, database schema, API development, UI components, and documentation. The team agreed on a Git workflow using feature branches with pull request reviews before merging.",
     "Key concerns included SQLite concurrency limitations (deemed acceptable for demo scope) and open questions around taboo word filtering edge cases that needed clarification."),
    ("Meeting 2 \u2014 Late March 2026",
     "The team completed and submitted the Phase I Software Requirements Specification (Version 1.0), covering all use cases, supplementary requirements, and system assumptions. The database schema design was reviewed and finalized with 11 models. A working login page prototype was demonstrated, validating the chosen tech stack.",
     "Concerns were raised about API key management for the AI chat feature, with the team agreeing to use environment variables and exclude keys from version control. The tight timeline for Phase II was noted and pseudo-code assignments began."),
    ("Meeting 3 \u2014 Early April 2026",
     "Significant implementation progress was reviewed. The database schema with all 11 models and a comprehensive seed script (16 users across all roles) had been merged to the main branch. The authentication system, cookie-based session management, and role-specific dashboards (student, instructor, registrar) were completed and demonstrated. The team reviewed the integrated portal UI with the midnight navy theme.",
     "Remaining work was discussed, including course registration, grading, reviews, complaint processing, and graduation workflows. The team coordinated responsibilities for the next sprint."),
    ("Meeting 4 \u2014 Mid-April 2026",
     "The team began work on the Phase II Design Report. Responsibilities were divided across the team for the E-R diagram, pseudo-code, sequence diagrams, and Petri-net diagrams. All use-case scenarios were reviewed for completeness and consistency with the Phase I SRS. The team also discussed the remaining implementation roadmap for Phase III.",
     "The primary concern was the approaching Phase II deadline. The team also discussed contingency plans for the AI chat feature in case of API key issues."),
]

for title, body, concerns in meetings:
    heading(title, 2)
    para(body)
    para("Concerns:", bold=True)
    para(concerns)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. REPOSITORY
# ══════════════════════════════════════════════════════════════
heading("7. Repository", 1)
para("GitHub Repository: https://github.com/dreyesliranzo/CunyZeroLite", bold=True)

para("Branch structure:")
for b in [
    "main \u2014 stable, reviewed code only",
    "feature/database-schema \u2014 Prisma schema and seed data (merged)",
    "feature/auth-dashboards \u2014 Authentication, session management, role-based dashboards",
    "feature/student-portal \u2014 Student portal UI (merged)",
    "feature/ai-chat \u2014 AI assistant integration (in progress)",
]: bullet(b)

para("Technology Stack:")
for t in [
    "Next.js 16 (React framework)",
    "TypeScript (type safety)",
    "Prisma 7.5 (ORM)",
    "SQLite via better-sqlite3 (database)",
    "Tailwind CSS 4 (styling)",
    "OpenAI GPT-4o-mini (AI assistant)",
]: bullet(t)

# ── Header/Footer ──
add_header_footer()

# ── SAVE ──
output = "/home/samia/CunyZeroLite/docs/GroupK_CunyZero_PhaseII.docx"
doc.save(output)
print(f"Saved to {output}")
